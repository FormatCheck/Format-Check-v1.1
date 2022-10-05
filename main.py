# import modules
from docx import Document  # for accessing the document by python-docx
import docx2txt as docx  # for accessing the document by docx2txt
import streamlit as st  # for web app
from PIL import Image  # for logo
import re  # for regex pattern used in reference counter

# configure web app logo and name
logo = Image.open('logo.png')
st.set_page_config(page_title='Format Check v1.1', page_icon=logo, layout='wide')


# font name program function
def font_name():
    # TODO start of Font Name code --------------------
    # add font name program banner
    st.subheader("**Font Name**")
    # access Word document file
    WordFile = Document(uploaded_file)

    # TODO Font Name code for TITLE ENGG1500 style
    # check font name for TITLE ENGG1500 style (text on title page) ********************
    title_font = set()  # store all TITLE ENGG1500 style font names in the set title_font
    title_wrong_font = set()  # store unacceptable TITLE ENGG1500 style font names in the set title_wrong_font
    title_wrong_font_words = sorted(set())  # store TITLE ENGG1500 style text that are in unacceptable fonts in the sorted list title_wrong_font_words
    CORRECT_FONT_NAME_TITLE = None  # state the specified font for TITLE ENGG1500 style and store in the variable CORRECT_FONT_NAME_TITLE
    TITLE_TEXT_STYLE = 'TITLE ENGG1500'  # state the specified style name for title page text and store in the variable TITLE_TEXT_STYLE
    for paragraph in WordFile.paragraphs:
        if TITLE_TEXT_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add fonts from each run into the set title_font
                title_font.add(run.font.name)
                # check if fonts are unacceptable, if so, store in the set title_wrong_font
                if run.font.name is not CORRECT_FONT_NAME_TITLE:
                    title_wrong_font.add(run.font.name)
                    # append TITLE ENGG1500 style text that contain unacceptable fonts in the sorted list title_wrong_font_words
                    title_wrong_font_words.append(run.text)

    # check if any elements in title_font are not CORRECT_FONT_NAME_TITLE and if title_font is not empty
    if any(name is not CORRECT_FONT_NAME_TITLE for name in title_font) and len(title_font) != 0:
        # print this if any elements in title_font are not CORRECT_FONT_NAME_TITLE and if title_font is not empty and print title_wrong_font and title_wrong_font_words contents
        st.error(f'''
        ❌ Text written in the TITLE ENGG1500 style have incorrect font(s): {', '.join(map(str, title_wrong_font))}  
        🡆 Incorrect font(s) found here: {' >> '.join(map(str, title_wrong_font_words))}
        ''')
    # check if title_font is empty, if so this means that TITLE ENGG1500 style was not found
    elif len(title_font) == 0:
        # print this if title_font is empty, since TITLE ENGG1500 style was not found
        st.info("ℹ️TITLE ENGG1500 style font name not found as this style was not used.")

    # TODO Font Name code for HEADING 1 ENGG1500 style
    # check font name for HEADING 1 ENGG1500 style (main headings) ********************
    h1_font = set()  # store all HEADING 1 ENGG1500 style font names in the set h1_font
    h1_wrong_font = set()  # store unacceptable HEADING 1 ENGG1500 style font names in the set h1_wrong_font
    h1_wrong_font_words = sorted(set())  # store HEADING 1 ENGG1500 style text that are in unacceptable fonts in the sorted list h1_wrong_font_words
    CORRECT_FONT_NAME_H1 = None  # state the specified font for HEADING 1 ENGG1500 style and store in the variable CORRECT_FONT_NAME_H1
    H1_STYLE = 'HEADING 1 ENGG1500'  # state the specified style name for main headings and store in the variable H1_STYLE
    for paragraph in WordFile.paragraphs:
        if H1_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add fonts from each run into the set h1_font
                h1_font.add(run.font.name)
                # check if fonts are unacceptable, if so, store in the set h1_wrong_font
                if run.font.name is not CORRECT_FONT_NAME_H1:
                    h1_wrong_font.add(run.font.name)
                    # append HEADING 1 ENGG1500 style text that contain unacceptable fonts in the sorted list h1_wrong_font_words
                    h1_wrong_font_words.append(run.text)

    # check if any elements in h1_font are not CORRECT_FONT_NAME_H1 and if h1_font is not empty
    if any(name is not CORRECT_FONT_NAME_H1 for name in h1_font) and len(h1_font) != 0:
        # print this if any elements in h1_font are not CORRECT_FONT_NAME_H1 and if h1_font is not empty and print h1_wrong_font and h1_wrong_font_words contents
        st.error(f'''
        ❌ Text written in the HEADING 1 ENGG1500 style have incorrect font(s): {', '.join(map(str, h1_wrong_font))}  
        🡆 Incorrect font(s) found here: {' >> '.join(map(str, h1_wrong_font_words))}
        ''')
    # check if h1_font is empty, if so this means that HEADING 1 ENGG1500 style was not found
    elif len(h1_font) == 0:
        # print this if h1_font is empty, since HEADING 1 ENGG1500 style was not found
        st.info("ℹ️HEADING 1 ENGG1500 style font name not found as this style was not used.")

    # TODO Font Name code for HEADING 2 ENGG1500 style
    # check font name for HEADING 2 ENGG1500 style (sub headings) ********************
    h2_font = set()  # store all HEADING 2 ENGG1500 style font names in the set h2_font
    h2_wrong_font = set()  # store unacceptable HEADING 2 ENGG1500 style font names in the set h2_wrong_font
    h2_wrong_font_words = sorted(set())  # store HEADING 2 ENGG1500 style text that are in unacceptable fonts in the sorted list h2_wrong_font_words
    CORRECT_FONT_NAME_H2 = None  # state the specified font for HEADING 2 ENGG1500 style and store in the variable CORRECT_FONT_NAME_H2
    H2_STYLE = 'HEADING 2 ENGG1500'  # state the specified style name for sub headings and store in the variable H2_STYLE
    for paragraph in WordFile.paragraphs:
        if H2_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add fonts from each run into the set h2_font
                h2_font.add(run.font.name)
                # check if fonts are unacceptable, if so, store in the set h2_wrong_font
                if run.font.name is not CORRECT_FONT_NAME_H2:
                    h2_wrong_font.add(run.font.name)
                    # append HEADING 2 ENGG1500 style text that contain unacceptable fonts in the sorted list h2_wrong_font_words
                    h2_wrong_font_words.append(run.text)

    # check if any elements in h2_font are not CORRECT_FONT_NAME_H2 and if h2_font is not empty
    if any(name is not CORRECT_FONT_NAME_H2 for name in h2_font) and len(h2_font) != 0:
        # print this if any elements in h2_font are not CORRECT_FONT_NAME_H2 and if h2_font is not empty and print h2_wrong_font and h2_wrong_font_words contents
        st.error(f'''
        ❌ Text written in the HEADING 2 ENGG1500 style have incorrect font(s): {', '.join(map(str, h2_wrong_font))}  
        🡆 Incorrect font(s) found here: {' >> '.join(map(str, h2_wrong_font_words))}
        ''')
    # check if h2_font is empty, if so this means that HEADING 2 ENGG1500 style was not found
    elif len(h2_font) == 0:
        # print this if h2_font is empty, since HEADING 2 ENGG1500 style was not found
        st.info("ℹ️HEADING 2 ENGG1500 style font name not found as this style was not used.")

    # TODO Font Name code for HEADING 3 ENGG1500 style
    # check font name for HEADING 3 ENGG1500 style (sub headings) ********************
    h3_font = set()  # store all HEADING 3 ENGG1500 style font names in the set h3_font
    h3_wrong_font = set()  # store unacceptable HEADING 3 ENGG1500 style font names in the set h3_wrong_font
    h3_wrong_font_words = sorted(set())  # store HEADING 3 ENGG1500 style text that are in unacceptable fonts in the sorted list h3_wrong_font_words
    CORRECT_FONT_NAME_H3 = None  # state the specified font for HEADING 3 ENGG1500 style and store in the variable CORRECT_FONT_NAME_H3
    H3_STYLE = 'HEADING 3 ENGG1500'  # state the specified style name for sub headings and store in the variable H3_STYLE
    for paragraph in WordFile.paragraphs:
        if H3_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add fonts from each run into the set h3_font
                h3_font.add(run.font.name)
                # check if fonts are unacceptable, if so, store in the set h3_wrong_font
                if run.font.name is not CORRECT_FONT_NAME_H3:
                    h3_wrong_font.add(run.font.name)
                    # append HEADING 3 ENGG1500 style text that contain unacceptable fonts in the sorted list h3_wrong_font_words
                    h3_wrong_font_words.append(run.text)

    # check if any elements in h3_font are not CORRECT_FONT_NAME_H3 and if h3_font is not empty
    if any(name is not CORRECT_FONT_NAME_H3 for name in h3_font) and len(h3_font) != 0:
        # print this if any elements in h3_font are not CORRECT_FONT_NAME_H3 and if h3_font is not empty and print h3_wrong_font and h3_wrong_font_words contents
        st.error(f'''
        ❌ Text written in the HEADING 3 ENGG1500 style have incorrect font(s): {', '.join(map(str, h3_wrong_font))}  
        🡆 Incorrect font(s) found here: {' >> '.join(map(str, h3_wrong_font_words))}
        ''')
    # check if h3_font is empty, if so this means that HEADING 3 ENGG1500 style was not found
    elif len(h3_font) == 0:
        # print this if h3_font is empty, since HEADING 3 ENGG1500 style was not found
        st.info("ℹ️HEADING 3 ENGG1500 style font name not found as this style was not used.")

    # TODO Font Name code for HEADING 4 ENGG1500 style
    # check font name for HEADING 4 ENGG1500 style (sub headings) ********************
    h4_font = set()  # store all HEADING 4 ENGG1500 style font names in the set h4_font
    h4_wrong_font = set()  # store unacceptable HEADING 4 ENGG1500 style font names in the set h4_wrong_font
    h4_wrong_font_words = sorted(set())  # store HEADING 4 ENGG1500 style text that are in unacceptable fonts in the sorted list h4_wrong_font_words
    CORRECT_FONT_NAME_H4 = None  # state the specified font for HEADING 4 ENGG1500 style and store in the variable CORRECT_FONT_NAME_H4
    H4_STYLE = 'HEADING 4 ENGG1500'  # state the specified style name for sub headings and store in the variable H4_STYLE
    for paragraph in WordFile.paragraphs:
        if H4_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add fonts from each run into the set h4_font
                h4_font.add(run.font.name)
                # check if fonts are unacceptable, if so, store in the set h4_wrong_font
                if run.font.name is not CORRECT_FONT_NAME_H4:
                    h4_wrong_font.add(run.font.name)
                    # append HEADING 4 ENGG1500 style text that contain unacceptable fonts in the sorted list h4_wrong_font_words
                    h4_wrong_font_words.append(run.text)

    # check if any elements in h4_font are not CORRECT_FONT_NAME_H4 and if h4_font is not empty
    if any(name is not CORRECT_FONT_NAME_H4 for name in h4_font) and len(h4_font) != 0:
        # print this if any elements in h4_font are not CORRECT_FONT_NAME_H4 and if h4_font is not empty and print h4_wrong_font and h4_wrong_font_words contents
        st.error(f'''
        ❌ Text written in the HEADING 4 ENGG1500 style have incorrect font(s): {', '.join(map(str, h4_wrong_font))}  
        🡆 Incorrect font(s) found here: {' >> '.join(map(str, h4_wrong_font_words))}
        ''')
    # check if h4_font is empty, if so this means that HEADING 4 ENGG1500 style was not found
    elif len(h4_font) == 0:
        # print this if h4_font is empty, since HEADING 4 ENGG1500 style was not found
        st.info("ℹ️HEADING 4 ENGG1500 style font name not found as this style was not used.")

    # TODO Font Name code for BODY ENGG1500 style
    # check font name for BODY ENGG1500 style (paragraphs) ********************
    body_font = set()  # store all BODY ENGG1500 style font names in the set body_font
    body_wrong_font = set()  # store unacceptable BODY ENGG1500 style font names in the set body_wrong_font
    body_wrong_font_words = sorted(set())  # store BODY ENGG1500 style text that are in unacceptable fonts in the sorted list body_wrong_font_words
    CORRECT_FONT_NAME_BODY = None  # state the specified font for BODY ENGG1500 style and store in the variable CORRECT_FONT_NAME_NORM
    PARAGRAPH_STYLE = 'BODY ENGG1500'  # state the specified style name for paragraphs and store in the variable PARAGRAPH_STYLE
    for paragraphs in WordFile.paragraphs:
        if PARAGRAPH_STYLE == paragraphs.style.name:
            for run in paragraphs.runs:
                # add fonts from each run into the set body_font
                body_font.add(run.font.name)
                # check if fonts are unacceptable, if so, store in the set body_wrong_font
                if run.font.name is not CORRECT_FONT_NAME_BODY:
                    body_wrong_font.add(run.font.name)
                    # append BODY ENGG1500 style text that contain unacceptable fonts in the sorted list body_wrong_font_words
                    body_wrong_font_words.append(run.text)

    # check if any elements in body_font are not CORRECT_FONT_NAME_BODY
    if any(name is not CORRECT_FONT_NAME_BODY for name in body_font) and len(body_font) != 0:
        # print this if any elements in body_font are not CORRECT_FONT_NAME_BODY and print body_wrong_font and body_wrong_font_words contents
        st.error(f'''
        ❌ Text written in the BODY ENGG1500 style have incorrect font(s): {', '.join(map(str, body_wrong_font))}  
        🡆 Incorrect font(s) found here: {' >> '.join(map(str, body_wrong_font_words))}
        ''')
    # check if body_font is empty, if so this means that BODY ENGG1500 style was not found
    elif len(body_font) == 0:
        # print this if body_font is empty, which means that BODY ENGG1500 style was not found
        st.info("ℹ️BODY ENGG1500 style font name not found as this style was not used.")

    # TODO Font Name code for BULLET LIST ENGG1500 style
    # check font name for BULLET LIST ENGG1500 style (dot points) ********************
    bullet_font = set()  # store all BULLET LIST ENGG1500 style font names in the set bullet_font
    bullet_wrong_font = set()  # store unacceptable BULLET LIST ENGG1500 style font names in the set bullet_wrong_font
    bullet_wrong_font_words = sorted(set())  # store BULLET LIST ENGG1500 style text that are in unacceptable fonts in the sorted list bullet_wrong_font_words
    CORRECT_FONT_NAME_BULLET = None  # state the specified font for BULLET LIST ENGG1500 style and store in the variable CORRECT_FONT_NAME_LIST
    DOT_POINT_STYLE = 'BULLET LIST ENGG1500'  # state the specified style name for dot points and store in the variable DOT_POINT_STYLE
    for paragraph in WordFile.paragraphs:
        if DOT_POINT_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add fonts from each run into the set bullet_font
                bullet_font.add(run.font.name)
                # check if fonts are unacceptable, if so, store in the set bullet_wrong_font
                if run.font.name is not CORRECT_FONT_NAME_BULLET:
                    bullet_wrong_font.add(run.font.name)
                    # append BULLET LIST ENGG1500 style text that contain unacceptable fonts in the sorted list bullet_wrong_font_words
                    bullet_wrong_font_words.append(run.text)

    # check if any elements in bullet_font are not CORRECT_FONT_NAME_BULLET and if bullet_font is not empty
    if any(name is not CORRECT_FONT_NAME_BULLET for name in bullet_font) and len(bullet_font) != 0:
        # print this if any elements in bullet_font are not CORRECT_FONT_NAME_BULLET and if bullet_font is not empty and print bullet_wrong_font and bullet_wrong_font_words contents
        st.error(f'''
        ❌ Text written in the BULLET LIST ENGG1500 style have incorrect font(s): {', '.join(map(str, bullet_wrong_font))}  
        🡆 Incorrect font(s) found here: {' >> '.join(map(str, bullet_wrong_font_words))}
        ''')
    # check if bullet_font is empty, if so this means that BULLET LIST ENGG1500 style was not found
    elif len(bullet_font) == 0:
        # print this if bullet_font is empty, since BULLET LIST ENGG1500 style was not found
        st.info("ℹ️BULLET LIST ENGG1500 style font name not found as this style was not used.")

    # TODO Font Name code for CAPTION ENGG1500 style
    # check font name for CAPTION ENGG1500 style (figure and table captions) ********************
    cap_font = set()  # store all CAPTION ENGG1500 style font names in the set cap_font
    cap_wrong_font = set()  # store unacceptable CAPTION ENGG1500 style font names in the set cap_wrong_font
    cap_wrong_font_words = sorted(set())  # store CAPTION ENGG1500 style text that are in unacceptable fonts in the sorted list cap_wrong_font_words
    CORRECT_FONT_NAME_CAP = None  # state the specified font for CAPTION ENGG1500 style and store in the variable CORRECT_FONT_NAME_CAP
    CAPTION_STYLE = 'CAPTION ENGG1500'  # state the specified style name for figure and table captions and store in the variable CAPTION_STYLE
    for paragraph in WordFile.paragraphs:
        if CAPTION_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add fonts from each run into the set cap_font
                cap_font.add(run.font.name)
                # check if fonts are unacceptable, if so, store in the set cap_wrong_font
                if run.font.name is not CORRECT_FONT_NAME_CAP:
                    cap_wrong_font.add(run.font.name)
                    # append CAPTION ENGG1500 style text that contain unacceptable fonts in the sorted list norm_wrong_font_words
                    cap_wrong_font_words.append(run.text)

    # check if any elements in cap_font are not CORRECT_FONT_NAME_CAP and if cap_font is not empty
    if any(name is not CORRECT_FONT_NAME_CAP for name in cap_font) and len(cap_font) != 0:
        # print this if any elements in cap_font are not CORRECT_FONT_NAME_CAP and if cap_font is not empty and print cap_wrong_font and cap_wrong_font_words contents
        st.error(f'''
        ❌ Text written in the CAPTION ENGG1500 style have incorrect font(s): {', '.join(map(str, cap_wrong_font))}  
        🡆 Incorrect font(s) found here: {' >> '.join(map(str, cap_wrong_font_words))}
        ''')
    # check if cap_font is empty, if so this means that CAPTION ENGG1500 style was not found
    elif len(cap_font) == 0:
        # print this if cap_font is empty, since CAPTION ENGG1500 style was not found
        st.info("ℹ️CAPTION ENGG1500 style font name not found as this style was not used.")

    # TODO Font Name code for TABLE ENGG1500 style
    # check font name for TABLE ENGG1500 style (text in tables) ********************
    table_font = set()  # store all TABLE ENGG1500 style font names in the set table_font
    table_wrong_font = set()  # store unacceptable TABLE ENGG1500 style font names in the set table_wrong_font
    table_wrong_font_words = sorted(set())  # store TABLE ENGG1500 style text that are in unacceptable fonts in the sorted list table_wrong_font_words
    CORRECT_FONT_NAME_TABLE = None  # state the specified font for TABLE ENGG1500 style and store in the variable CORRECT_FONT_NAME_TABLE
    TABLE_TEXT_STYLE = 'TABLE ENGG1500'  # state the specified style name for text in tables and store in the variable TABLE_TEXT_STYLE
    for table in WordFile.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if TABLE_TEXT_STYLE == paragraph.style.name:
                        for run in paragraph.runs:
                            # add fonts from each run into the set table_font
                            table_font.add(run.font.name)
                            # check if fonts are unacceptable, if so, store in the set table_wrong_font
                            if run.font.name is not CORRECT_FONT_NAME_TABLE:
                                table_wrong_font.add(run.font.name)
                                # append TABLE ENGG1500 style text that contain unacceptable fonts in the sorted list table_wrong_font_words
                                table_wrong_font_words.append(run.text)

    # check if any elements in table_font are not CORRECT_FONT_NAME_TABLE
    if any(name is not CORRECT_FONT_NAME_TABLE for name in table_font) and len(table_font) != 0:
        # print this if any elements in table_font are not CORRECT_FONT_NAME_TABLE and print table_wrong_font and table_wrong_font_words contents
        st.error(f'''
        ❌ Text written in the TABLE ENGG1500 style have incorrect font(s): {', '.join(map(str, table_wrong_font))}  
        🡆 Incorrect font(s) found here: {' >> '.join(map(str, table_wrong_font_words))}
        ''')
    # check if table_font is empty, if so this means that TABLE ENGG1500 style was not found
    elif len(table_font) == 0:
        # print this if table_font is empty, which means that TABLE ENGG1500 style was not found
        st.info("ℹ️TABLE ENGG1500 style font name not found as this style was not used.")


# font size program function
def font_size():
    # TODO start of Font Size code --------------------
    # add font size program banner
    st.subheader("**Font Size**")
    # access Word document file
    WordFile = Document(uploaded_file)

    # TODO Font Size code for TITLE ENGG1500 style
    # check font size for TITLE ENGG1500 style (text on title page) ********************
    title_size = set()  # store all TITLE ENGG1500 style font sizes in the set title_size
    title_wrong_size = set()  # store unacceptable TITLE ENGG1500 style font sizes in the set title_wrong_size
    title_wrong_size_words = sorted(set())  # store TITLE ENGG1500 style text that are in unacceptable font size in the sorted list title_wrong_size_words
    CORRECT_FONT_SIZE_TITLE = None  # state the specified font size for TITLE ENGG1500 style and store in the variable CORRECT_FONT_SIZE_TITLE
    TITLE_TEXT_STYLE = 'TITLE ENGG1500'  # state the specified style name for title page text and store in the variable TITLE_TEXT_STYLE
    for paragraph in WordFile.paragraphs:
        if TITLE_TEXT_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add font sizes from each run into the set title_size
                title_size.add(run.font.size)
                # check if font sizes are unacceptable, if so, store in the set title_wrong_size
                if run.font.size is not CORRECT_FONT_SIZE_TITLE:
                    title_wrong_size.add(run.font.size/12700)
                    # append TITLE ENGG1500 style text that contain unacceptable font sizes in the sorted list title_wrong_size_words
                    title_wrong_size_words.append(run.text)

    # check if any elements in title_size are not CORRECT_FONT_SIZE_TITLE and if title_size is not empty
    if any(size is not CORRECT_FONT_SIZE_TITLE for size in title_size) and len(title_size) != 0:
        # print this if any elements in title_size are not CORRECT_FONT_SIZE_TITLE and if title_size is not empty and print title_wrong_size and title_wrong_size_words contents
        st.error(f'''
        ❌ Text written in the TITLE ENGG1500 style have incorrect font size(s): {', '.join(map(str, title_wrong_size))}  
        🡆 Incorrect font size(s) found here: {' >> '.join(map(str, title_wrong_size_words))}
        ''')
    # check if title_size is empty, if so this means that TITLE ENGG1500 style was not found
    elif len(title_size) == 0:
        # print this if title_size is empty, since TITLE ENGG1500 style was not found
        st.info("ℹ️TITLE ENGG1500 style font size not found as this style was not used.")

    # TODO Font Size code for HEADING 1 ENGG1500 style
    # check font size for HEADING 1 ENGG1500 style (main headings) ********************
    h1_size = set()  # store all HEADING 1 ENGG1500 style font sizes in the set h1_size
    h1_wrong_size = set()  # store unacceptable HEADING 1 ENGG1500 style font sizes in the set h1_wrong_size
    h1_wrong_size_words = sorted(set())  # store HEADING 1 ENGG1500 style text that are in unacceptable font size in the sorted list h1_wrong_size_words
    CORRECT_FONT_SIZE_H1 = None  # state the specified font size for HEADING 1 ENGG1500 style and store in the variable CORRECT_FONT_SIZE_H1
    H1_STYLE = 'HEADING 1 ENGG1500'  # state the specified style name for main headings and store in the variable H1_STYLE
    for paragraph in WordFile.paragraphs:
        if H1_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add font sizes from each run into the set h1_size
                h1_size.add(run.font.size)
                # check if font sizes are unacceptable, if so, store in the set h1_wrong_size
                if run.font.size is not CORRECT_FONT_SIZE_H1:
                    h1_wrong_size.add(run.font.size/12700)
                    # append HEADING 1 ENGG1500 style text that contain unacceptable font sizes in the sorted list h1_wrong_size_words
                    h1_wrong_size_words.append(run.text)

    # check if any elements in h1_size are not CORRECT_FONT_SIZE_H1 and if h1_size is not empty
    if any(size is not CORRECT_FONT_SIZE_H1 for size in h1_size) and len(h1_size) != 0:
        # print this if any elements in h1_size are not CORRECT_FONT_SIZE_H1 and if h1_size is not empty and print h1_wrong_size and h1_wrong_size_words contents
        st.error(f'''
        ❌ Text written in the HEADING 1 ENGG1500 style have incorrect font size(s): {', '.join(map(str, h1_wrong_size))}  
        🡆 Incorrect font size(s) found here: {' >> '.join(map(str, h1_wrong_size_words))}
        ''')
    # check if h1_size is empty, if so this means that HEADING 1 ENGG1500 style was not found
    elif len(h1_size) == 0:
        # print this if h1_size is empty, since HEADING 1 ENGG1500 style was not found
        st.info("ℹ️HEADING 1 ENGG1500 style font size not found as this style was not used.")

    # TODO Font Size code for HEADING 2 ENGG1500 style
    # check font size for HEADING 2 ENGG1500 style (sub headings) ********************
    h2_size = set()  # store all HEADING 2 ENGG1500 style font sizes in the set h2_size
    h2_wrong_size = set()  # store unacceptable HEADING 2 ENGG1500 style font sizes in the set h2_wrong_size
    h2_wrong_size_words = sorted(set())  # store HEADING 2 ENGG1500 style text that are in unacceptable font size in the sorted list h2_wrong_size_words
    CORRECT_FONT_SIZE_H2 = None  # state the specified font size for HEADING 2 ENGG1500 style and store in the variable CORRECT_FONT_SIZE_H2
    H2_STYLE = 'HEADING 2 ENGG1500'  # state the specified style name for sub headings and store in the variable H2_STYLE
    for paragraph in WordFile.paragraphs:
        if H2_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add font sizes from each run into the set h2_size
                h2_size.add(run.font.size)
                # check if font sizes are unacceptable, if so, store in the set h2_wrong_size
                if run.font.size is not CORRECT_FONT_SIZE_H2:
                    h2_wrong_size.add(run.font.size/12700)
                    # append HEADING 2 ENGG1500 style text that contain unacceptable font sizes in the sorted list h2_wrong_size_words
                    h2_wrong_size_words.append(run.text)

    # check if any elements in h2_size are not CORRECT_FONT_SIZE_H2 and if h2_size is not empty
    if any(size is not CORRECT_FONT_SIZE_H2 for size in h2_size) and len(h2_size) != 0:
        # print this if any elements in h2_size are not CORRECT_FONT_SIZE_H2 and if h2_size is not empty and print h2_wrong_size and h2_wrong_size_words contents
        st.error(f'''
        ❌ Text written in the HEADING 2 ENGG1500 style have incorrect font size(s): {', '.join(map(str, h2_wrong_size))}  
        🡆 Incorrect font size(s) found here: {' >> '.join(map(str, h2_wrong_size_words))}
        ''')
    # check if h2_size is empty, if so this means that HEADING 2 ENGG1500 style was not found
    elif len(h2_size) == 0:
        # print this if h2_size is empty, since HEADING 2 ENGG1500 style was not found
        st.info("ℹ️HEADING 2 ENGG1500 style font size not found as this style was not used.")

    # TODO Font Size code for HEADING 3 ENGG1500 style
    # check font size for HEADING 3 ENGG1500 style (sub headings) ********************
    h3_size = set()  # store all HEADING 3 ENGG1500 style font sizes in the set h3_size
    h3_wrong_size = set()  # store unacceptable HEADING 3 ENGG1500 style font sizes in the set h3_wrong_size
    h3_wrong_size_words = sorted(set())  # store HEADING 3 ENGG1500 style text that are in unacceptable font size in the sorted list h3_wrong_size_words
    CORRECT_FONT_SIZE_H3 = None  # state the specified font size for HEADING 3 ENGG1500 style and store in the variable CORRECT_FONT_SIZE_H3
    H3_STYLE = 'HEADING 3 ENGG1500'  # state the specified style name for sub headings and store in the variable H3_STYLE
    for paragraph in WordFile.paragraphs:
        if H3_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add font sizes from each run into the set h3_size
                h3_size.add(run.font.size)
                # check if font sizes are unacceptable, if so, store in the set h3_wrong_size
                if run.font.size is not CORRECT_FONT_SIZE_H3:
                    h3_wrong_size.add(run.font.size/12700)
                    # append HEADING 3 ENGG1500 style text that contain unacceptable font sizes in the sorted list h3_wrong_size_words
                    h3_wrong_size_words.append(run.text)

    # check if any elements in h3_size are not CORRECT_FONT_SIZE_H3 and if h3_size is not empty
    if any(size is not CORRECT_FONT_SIZE_H3 for size in h3_size) and len(h3_size) != 0:
        # print this if any elements in h3_size are not CORRECT_FONT_SIZE_H3 and if h3_size is not empty and print h3_wrong_size and h3_wrong_size_words contents
        st.error(f'''
        ❌ Text written in the HEADING 3 ENGG1500 style have incorrect font size(s): {', '.join(map(str, h3_wrong_size))}  
        🡆 Incorrect font size(s) found here: {' >> '.join(map(str, h3_wrong_size_words))}
        ''')
    # check if h3_size is empty, if so this means that HEADING 3 ENGG1500 style was not found
    elif len(h3_size) == 0:
        # print this if h3_size is empty, since HEADING 3 ENGG1500 style was not found
        st.info("ℹ️HEADING 3 ENGG1500 style font size not found as this style was not used.")

    # TODO Font Size code for HEADING 4 ENGG1500 style
    # check font size for HEADING 4 ENGG1500 style (sub headings) ********************
    h4_size = set()  # store all HEADING 4 ENGG1500 style font sizes in the set h4_size
    h4_wrong_size = set()  # store unacceptable HEADING 4 ENGG1500 style font sizes in the set h4_wrong_size
    h4_wrong_size_words = sorted(set())  # store HEADING 4 ENGG1500 style text that are in unacceptable font size in the sorted list h4_wrong_size_words
    CORRECT_FONT_SIZE_H4 = None  # state the specified font size for HEADING 4 ENGG1500 style and store in the variable CORRECT_FONT_SIZE_H4
    H4_STYLE = 'HEADING 4 ENGG1500'  # state the specified style name for sub headings and store in the variable H4_STYLE
    for paragraph in WordFile.paragraphs:
        if H4_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add font sizes from each run into the set h4_sizes
                h4_size.add(run.font.size)
                # check if font sizes are unacceptable, if so, store in the set h4_wrong_size
                if run.font.size is not CORRECT_FONT_SIZE_H4:
                    h4_wrong_size.add(run.font.size/12700)
                    # append HEADING 4 ENGG1500 style text that contain unacceptable font sizes in the sorted list h4_wrong_size_words
                    h4_wrong_size_words.append(run.text)

    # check if any elements in h4_size are not CORRECT_FONT_SIZE_H4 and if h4_size is not empty
    if any(size is not CORRECT_FONT_SIZE_H4 for size in h4_size) and len(h4_size) != 0:
        # print this if any elements in h4_size are not CORRECT_FONT_SIZE_H4 and if h4_size is not empty and print h4_wrong_size and h4_wrong_size_words contents
        st.error(f'''
        ❌ Text written in the HEADING 4 ENGG1500 style have incorrect font size(s): {', '.join(map(str, h4_wrong_size))}  
        🡆 Incorrect font size(s) found here: {' >> '.join(map(str, h4_wrong_size_words))}
        ''')
    # check if h4_size is empty, if so this means that HEADING 4 ENGG1500 style was not found
    elif len(h4_size) == 0:
        # print this if h4_size is empty, since HEADING 4 ENGG1500 style was not found
        st.info("ℹ️HEADING 4 ENGG1500 style font size not found as this style was not used.")

    # TODO Font Size code for BODY ENGG1500 style
    # check font size for BODY ENGG1500 style (paragraphs) ********************
    body_size = set()  # store all BODY ENGG1500 style font sizes in the set body_size
    body_wrong_size = set()  # store unacceptable BODY ENGG1500 style font sizes in the set body_wrong_size
    body_wrong_size_words = sorted(set())  # store BODY ENGG1500 style text that are in unacceptable font size in the sorted list body_wrong_size_words
    CORRECT_FONT_SIZE_BODY = None  # state the specified font size for BODY ENGG1500 style and store in the variable CORRECT_FONT_SIZE_BODY
    PARAGRAPH_STYLE = 'BODY ENGG1500'  # state the specified style name for paragraphs and store in the variable PARAGRAPH_STYLE
    for paragraph in WordFile.paragraphs:
        if PARAGRAPH_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add font sizes from each run into the set body_size
                body_size.add(run.font.size)
                # check if font sizes are unacceptable, if so, store in the set body_wrong_size
                if run.font.size is not CORRECT_FONT_SIZE_BODY:
                    body_wrong_size.add(run.font.size/12700)
                    # append BODY ENGG1500 style text that contain unacceptable font sizes in the sorted list body_wrong_size_words
                    body_wrong_size_words.append(run.text)

    # check if any elements in body_size are not CORRECT_FONT_SIZE_BODY
    if any(size is not CORRECT_FONT_SIZE_BODY for size in body_size) and len(body_size) != 0:
        # print this if any elements in body_font are not CORRECT_FONT_SIZE_BODY and print body_wrong_size and body_wrong_size_words contents
        st.error(f'''
        ❌ Text written in the BODY ENGG1500 style have incorrect font size(s): {', '.join(map(str, body_wrong_size))}  
        🡆 Incorrect font size(s) found here: {' >> '.join(map(str, body_wrong_size_words))}
        ''')
    # check if body_size is empty, if so this means that BODY ENGG1500 style was not found
    elif len(body_size) == 0:
        # print this if body_size is empty, which means that BODY ENGG1500 style was not found
        st.info("ℹ️BODY ENGG1500 style font size not found as this style was not used.")

    # TODO Font Size code for BULLET LIST ENGG1500 style
    # check font size for BULLET LIST ENGG1500 style (bullet list) ********************
    bullet_size = set()  # store all BULLET LIST ENGG1500 style font sizes in the set bullet_size
    bullet_wrong_size = set()  # store unacceptable BULLET LIST ENGG1500 style font sizes in the set bullet_wrong_size
    bullet_wrong_size_words = sorted(set())  # store BULLET LIST ENGG1500 style text that are in unacceptable font size in the sorted list bullet_wrong_size_words
    CORRECT_FONT_SIZE_BULLET = None  # state the specified font size for BULLET LIST ENGG1500 style and store in the variable CORRECT_FONT_SIZE_BULLET
    DOT_POINT_STYLE = 'BULLET LIST ENGG1500'  # state the specified style name for dot points and store in the variable DOT_POINT_STYLE
    for paragraph in WordFile.paragraphs:
        if DOT_POINT_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add font sizes from each run into the set bullet_size
                bullet_size.add(run.font.size)
                # check if font sizes are unacceptable, if so, store in the set bullet_wrong_size
                if run.font.size is not CORRECT_FONT_SIZE_BULLET:
                    bullet_wrong_size.add(run.font.size/12700)
                    # append BULLET LIST ENGG1500 style text that contain unacceptable font sizes in the sorted list bullet_wrong_size_words
                    bullet_wrong_size_words.append(run.text)

    # check if any elements in bullet_size are not CORRECT_FONT_SIZE_BULLET and if bullet_size is not empty
    if any(size is not CORRECT_FONT_SIZE_BULLET for size in bullet_size) and len(bullet_size) != 0:
        # print this if any elements in bullet_size are not CORRECT_FONT_SIZE_BULLET and if bullet_size is not empty and print bullet_wrong_size and bullet_wrong_size_words contents
        st.error(f'''
        ❌ Text written in the BULLET LIST ENGG1500 style have incorrect font size(s): {', '.join(map(str, bullet_wrong_size))}  
        🡆 Incorrect font size(s) found here: {' >> '.join(map(str, bullet_wrong_size_words))}
        ''')
    # check if bullet_size is empty, if so this means that BULLET LIST ENGG1500 style was not found
    elif len(bullet_size) == 0:
        # print this if bullet_size is empty, since BULLET LIST ENGG1500 style was not found
        st.info("ℹ️BULLET LIST ENGG1500 style font size not found as this style was not used.")

    # TODO Font Size code for CAPTION ENGG1500 style
    # check font size for CAPTION ENGG1500 style (figure and table captions) ********************
    cap_size = set()  # store all CAPTION ENGG1500 style font sizes in the set cap_size
    cap_wrong_size = set()  # store unacceptable CAPTION ENGG1500 style font sizes in the set cap_wrong_size
    cap_wrong_size_words = sorted(set())  # store CAPTION ENGG1500 style text that are in unacceptable font size in the sorted list cap_wrong_size_words
    CORRECT_FONT_SIZE_CAP = None  # state the specified font size for CAPTION ENGG1500 style and store in the variable CORRECT_FONT_SIZE_CAP
    CAPTION_STYLE = 'CAPTION ENGG1500'  # state the specified style name for figure and table captions and store in the variable CAPTION_STYLE
    for paragraph in WordFile.paragraphs:
        if CAPTION_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add font sizes from each run into the set cap_size
                cap_size.add(run.font.size)
                # check if font sizes are unacceptable, if so, store in the set cap_wrong_size
                if run.font.size is not CORRECT_FONT_SIZE_CAP:
                    cap_wrong_size.add(run.font.size/12700)
                    # append CAPTION ENGG1500 style text that contain unacceptable font sizes in the sorted list cap_wrong_size_words
                    cap_wrong_size_words.append(run.text)

    # check if any elements in cap_size are not CORRECT_FONT_SIZE_CAP and if cap_size is not empty
    if any(size is not CORRECT_FONT_SIZE_CAP for size in cap_size) and len(cap_size) != 0:
        # print this if any elements in cap_size are not CORRECT_FONT_SIZE_CAP and if cap_size is not empty and print cap_wrong_font and cap_wrong_font_words contents
        st.error(f'''
        ❌ Text written in the CAPTION ENGG1500 style have incorrect font size(s): {', '.join(map(str, cap_wrong_size))}  
        🡆 Incorrect font size(s) found here: {' >> '.join(map(str, cap_wrong_size_words))}
        ''')
    # check if cap_size is empty, if so this means that CAPTION ENGG1500 style was not found
    elif len(cap_size) == 0:
        # print this if cap_size is empty, since CAPTION ENGG1500 style was not found
        st.info("ℹ️CAPTION ENGG1500 style font size not found as this style was not used.")

    # TODO Font Size code for TABLE ENGG1500 style
    # check font size for TABLE ENGG1500 style text (text in tables) ********************
    table_size = set()  # store all TABLE ENGG1500 style font sizes in the set table_size
    table_wrong_size = set()  # store unacceptable TABLE ENGG1500 style font sizes in the set table_wrong_size
    table_wrong_size_words = sorted(set())  # store TABLE ENGG1500 style text that are in unacceptable font size in the sorted list table_wrong_size_words
    CORRECT_FONT_SIZE_TABLE = None  # state the specified font size for TABLE ENGG1500 style and store in the variable CORRECT_FONT_SIZE_TABLE
    TABLE_TEXT_STYLE = 'TABLE ENGG1500'  # state the specified style name for text in tables and store in the variable TABLE_TEXT_STYLE
    for table in WordFile.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if TABLE_TEXT_STYLE == paragraph.style.name:
                        for run in paragraph.runs:
                            # add font sizes from each run into the set table_size
                            table_size.add(run.font.size)
                            # check if fonts are unacceptable, if so, store in the set table_wrong_size
                            if run.font.size is not CORRECT_FONT_SIZE_TABLE:
                                table_wrong_size.add(run.font.size/12700)
                                # append TABLE ENGG1500 style text that contain unacceptable fonts in the sorted list table_wrong_size_words
                                table_wrong_size_words.append(run.text)

    # check if any elements in table_size are not CORRECT_FONT_SIZE_TABLE
    if any(size is not CORRECT_FONT_SIZE_TABLE for size in table_size) and len(table_size) != 0:
        # print this if any elements in table_size are not CORRECT_FONT_SIZE_TABLE and print table_wrong_size and table_wrong_size_words contents
        st.error(f'''
        ❌ Text written in the TABLE ENGG1500 style have incorrect font size(s): {', '.join(map(str, table_wrong_size))}  
        🡆 Incorrect font size(s) found here: {' >> '.join(map(str, table_wrong_size_words))}
        ''')
    # check if table_size is empty, if so this means that TABLE ENGG1500 style was not found
    elif len(table_size) == 0:
        # print this if table_size is empty, which means that TABLE ENGG1500 style was not found
        st.info("ℹ️TABLE ENGG1500 style font size not found as this style was not used.")


# font colour program function
def font_colour():
    # TODO start of Font Colour code --------------------
    # add font colour program banner
    st.subheader("**Font Colour**")
    # access Word document file
    WordFile = Document(uploaded_file)

    # TODO Font Colour code for TITLE ENGG1500 style
    # check font colour for TITLE ENGG1500 style (text on title page) ********************
    title_colour = set()  # store all TITLE ENGG1500 style font colours in the set title_colour
    title_wrong_colour = set()  # store unacceptable TITLE ENGG1500 style font colours in the set title_wrong_colour
    title_wrong_colour_words = sorted(set())  # store TITLE ENGG1500 style text that are in unacceptable font colour in the sorted list title_wrong_colour_words
    CORRECT_FONT_COLOUR_TITLE = None  # state the specified font colour for TITLE ENGG1500 style and store in the variable CORRECT_FONT_COLOUR_TITLE
    TITLE_TEXT_STYLE = 'TITLE ENGG1500'  # state the specified style name for title page text and store in the variable TITLE_TEXT_STYLE
    for paragraph in WordFile.paragraphs:
        if TITLE_TEXT_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add font colours from each run into the set title_colour
                title_colour.add(run.font.color.rgb)
                # check if font colours are unacceptable, if so, store in the set title_wrong_colour
                if run.font.color.rgb is not CORRECT_FONT_COLOUR_TITLE:
                    title_wrong_colour.add(run.font.color.rgb)
                    # append TITLE ENGG1500 style text that contain unacceptable font colours in the sorted list title_wrong_colour_words
                    title_wrong_colour_words.append(run.text)

    # check if any elements in title_colour are not CORRECT_FONT_COLOUR_TITLE and if title_colour is not empty
    if any(colour is not CORRECT_FONT_COLOUR_TITLE for colour in title_colour) and len(title_colour) != 0:
        # print this if any elements in title_colour are not CORRECT_FONT_COLOUR_TITLE and print title_wrong_colour and title_wrong_colour_words content
        st.error(f'''
        ❌ Text written in the TITLE ENGG1500 style have incorrect font colour(s): {', '.join(map(str, title_wrong_colour))}  
        🡆 Incorrect font colour(s) found here: {' >> '.join(map(str, title_wrong_colour_words))}
        ''')
    # check if title_colour is empty, if so this means that TITLE ENGG1500 style was not found
    elif len(title_colour) == 0:
        # print this if title_colour is empty, which means that TITLE ENGG1500 style was not found
        st.info("ℹ️TITLE ENGG1500 style font colour not found as this style was not used.")

    # TODO Font Colour code for HEADING 1 ENGG1500 style
    # check font colour for HEADING 1 ENGG1500 style (main headings) ********************
    h1_colour = set()  # store all HEADING 1 ENGG1500 style font colours in the set h1_colour
    h1_wrong_colour = set()  # store unacceptable HEADING 1 ENGG1500 style font colours in the set h1_wrong_colour
    h1_wrong_colour_words = sorted(set())  # store HEADING 1 ENGG1500 style text that are in unacceptable font colour in the sorted list h1_wrong_colour_words
    CORRECT_FONT_COLOUR_H1 = None  # state the specified font colour for HEADING 1 ENGG1500 style and store in the variable CORRECT_FONT_COLOUR_H1
    H1_STYLE = 'HEADING 1 ENGG1500'  # state the specified style name for main headings and store in the variable H1_STYLE
    for paragraph in WordFile.paragraphs:
        if H1_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add font colours from each run into the set h1_colour
                h1_colour.add(run.font.color.rgb)
                # check if font colours are unacceptable, if so, store in the set h1_wrong_colour
                if run.font.color.rgb is not CORRECT_FONT_COLOUR_H1:
                    h1_wrong_colour.add(run.font.color.rgb)
                    # append HEADING 1 ENGG1500 style text that contain unacceptable font colours in the sorted list h1_wrong_colour_words
                    h1_wrong_colour_words.append(run.text)

    # check if any elements in h1_colour are not CORRECT_FONT_COLOUR_H1 and if h1_colour is not empty
    if any(colour is not CORRECT_FONT_COLOUR_H1 for colour in h1_colour) and len(h1_colour) != 0:
        # print this if any elements in h1_colour are not CORRECT_FONT_COLOUR_H1 and print h1_wrong_colour and h1_wrong_colour_words content
        st.error(f'''
        ❌ Text written in the HEADING 1 ENGG1500 style have incorrect font colour(s): {', '.join(map(str, h1_wrong_colour))}  
        🡆 Incorrect font colours(s) found here: {' >> '.join(map(str, h1_wrong_colour_words))}
        ''')
    # check if h1_colour is empty, if so this means that HEADING 1 ENGG1500 style was not found
    elif len(h1_colour) == 0:
        # print this if h1_colour is empty, if so, this means that HEADING 1 ENGG1500 style was not found
        st.info("ℹ️HEADING 1 ENGG1500 style font colour not found as this style was not used.")

    # TODO Font Colour code for HEADING 2 ENGG1500 style
    # check font colour for HEADING 2 ENGG1500 style (sub headings) ********************
    h2_colour = set()  # store all HEADING 2 ENGG1500 style font colours in the set h2_colour
    h2_wrong_colour = set()  # store unacceptable HEADING 2 ENGG1500 style font colours in the set h2_wrong_colour
    h2_wrong_colour_words = sorted(set())  # store HEADING 2 ENGG1500 style text that are in unacceptable font colour in the sorted list h2_wrong_colour_words
    CORRECT_FONT_COLOUR_H2 = None  # state the specified font colour for HEADING 2 ENGG1500 style and store in the variable CORRECT_FONT_COLOUR_H2
    H2_STYLE = 'HEADING 2 ENGG1500'  # state the specified style name for sub headings and store in the variable H2_STYLE
    for paragraph in WordFile.paragraphs:
        if H2_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add font colours from each run into the set h2_colour
                h2_colour.add(run.font.color.rgb)
                # check if font colours are unacceptable, if so, store in the set h2_wrong_colour
                if run.font.color.rgb is not CORRECT_FONT_COLOUR_H2:
                    h2_wrong_colour.add(run.font.color.rgb)
                    # append HEADING 2 ENGG1500 style text that contain unacceptable font colours in the sorted list h2_wrong_colour_words
                    h2_wrong_colour_words.append(run.text)

    # check if any elements in h2_colour are not CORRECT_FONT_COLOUR_H2 and if h2_colour is not empty
    if any(colour is not CORRECT_FONT_COLOUR_H2 for colour in h2_colour) and len(h2_colour) != 0:
        # print this if any elements in h2_colour are not CORRECT_FONT_COLOUR_H2 and print h2_wrong_colour and h2_wrong_colour_words content
        st.error(f'''
        ❌ Text written in the HEADING 2 ENGG1500 style have incorrect font colour(s): {', '.join(map(str, h2_wrong_colour))}  
        🡆 Incorrect font colours(s) found here: {' >> '.join(map(str, h2_wrong_colour_words))}
        ''')
    # check if h2_colour is empty, if so this means that HEADING 2 ENGG1500 style was not found
    elif len(h2_colour) == 0:
        # print this if h2_colour is empty, which means that HEADING 2 ENGG1500 style was not found
        st.info("ℹ️HEADING 2 ENGG1500 style font colour not found as this style was not used.")

    # TODO Font Colour code for HEADING 3 ENGG1500 style
    # check font colour for HEADING 3 ENGG1500 style (sub headings) ********************
    h3_colour = set()  # store all HEADING 3 ENGG1500 style font colours in the set h3_colour
    h3_wrong_colour = set()  # store unacceptable HEADING 3 ENGG1500 style font colours in the set h3_wrong_colour
    h3_wrong_colour_words = sorted(set())  # store HEADING 3 ENGG1500 style text that are in unacceptable font colour in the sorted list h3_wrong_colour_words
    CORRECT_FONT_COLOUR_H3 = None  # state the specified font colour for HEADING 3 ENGG1500 style and store in the variable CORRECT_FONT_COLOUR_H3
    H3_STYLE = 'HEADING 3 ENGG1500'  # state the specified style name for sub headings and store in the variable H3_STYLE
    for paragraph in WordFile.paragraphs:
        if H3_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add font colours from each run into the set h3_colour
                h3_colour.add(run.font.color.rgb)
                # check if font colours are unacceptable, if so, store in the set h3_wrong_colour
                if run.font.color.rgb is not CORRECT_FONT_COLOUR_H3:
                    h3_wrong_colour.add(run.font.color.rgb)
                    # append HEADING 3 ENGG1500 style text that contain unacceptable font colours in the sorted list h3_wrong_colour_words
                    h3_wrong_colour_words.append(run.text)

    # check if any elements in h3_colour are not CORRECT_FONT_COLOUR_H3 and if h3_colour is not empty
    if any(colour is not CORRECT_FONT_COLOUR_H3 for colour in h3_colour) and len(h3_colour) != 0:
        # print this if any elements in h3_colour are not CORRECT_FONT_COLOUR_H3 and print h3_wrong_colour and h3_wrong_colour_words content
        st.error(f'''
        ❌ Text written in the HEADING 3 ENGG1500 style have incorrect font colour(s): {', '.join(map(str, h3_wrong_colour))}  
        🡆 Incorrect font colours(s) found here: {' >> '.join(map(str, h3_wrong_colour_words))}
        ''')
    # check if h3_colour is empty, if so this means that HEADING 3 ENGG1500 style was not found
    elif len(h3_colour) == 0:
        # print this if h3_colour is empty, which means that HEADING 3 ENGG1500 style was not found
        st.info("ℹ️HEADING 3 ENGG1500 style font colour not found as this style was not used.")

    # TODO Font Colour code for HEADING 4 ENGG1500 style
    # check font colour for HEADING 4 ENGG1500 style (sub headings) ********************
    h4_colour = set()  # store all HEADING 4 ENGG1500 style font colours in the set h4_colour
    h4_wrong_colour = set()  # store unacceptable HEADING 4 ENGG1500 style font colours in the set h4_wrong_colour
    h4_wrong_colour_words = sorted(set())  # store HEADING 4 ENGG1500 style text that are in unacceptable font colour in the sorted list h4_wrong_colour_words
    CORRECT_FONT_COLOUR_H4 = None  # state the specified font colour for HEADING 4 ENGG1500 style and store in the variable CORRECT_FONT_COLOUR_H4
    H4_STYLE = 'HEADING 4 ENGG1500'  # state the specified style name for sub headings and store in the variable H4_STYLE
    for paragraph in WordFile.paragraphs:
        if H4_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add font colours from each run into the set h4_colour
                h4_colour.add(run.font.color.rgb)
                # check if font colours are unacceptable, if so, store in the set h4_wrong_colour
                if run.font.color.rgb is not CORRECT_FONT_COLOUR_H4:
                    h4_wrong_colour.add(run.font.color.rgb)
                    # append HEADING 4 ENGG1500 style text that contain unacceptable font colours in the sorted list h4_wrong_colour_words
                    h4_wrong_colour_words.append(run.text)

    # check if any elements in h4_colour are not CORRECT_FONT_COLOUR_H4 and if h4_colour is not empty
    if any(colour is not CORRECT_FONT_COLOUR_H4 for colour in h4_colour) and len(h4_colour) != 0:
        # print this if any elements in h4_colour are not CORRECT_FONT_COLOUR_H4 and print h4_wrong_colour and h4_wrong_colour_words content
        st.error(f'''
        ❌ Text written in the HEADING 4 ENGG1500 style have incorrect font colour(s): {', '.join(map(str, h4_wrong_colour))}  
        🡆 Incorrect font colours(s) found here: {' >> '.join(map(str, h4_wrong_colour_words))}
        ''')
    # check if h4_colour is empty, if so this means that HEADING 4 ENGG1500 style was not found
    elif len(h4_colour) == 0:
        # print this if h4_colour is empty, which means that HEADING 4 ENGG1500 style was not found
        st.info("ℹ️HEADING 4 ENGG1500 style font colour not found as this style was not used.")

    # TODO Font Colour code for BODY ENGG1500 style
    # check font colour for BODY ENGG1500 style (paragraphs) ********************
    body_colour = set()  # store all BODY ENGG1500 style font colours in the set body_colour
    body_wrong_colour = set()  # store unacceptable BODY ENGG1500 style font colours in the set body_wrong_colour
    body_wrong_colour_words = sorted(set())  # store BODY ENGG1500 style text that are in unacceptable font colour in the sorted list body_wrong_colour_words
    CORRECT_FONT_COLOUR_BODY = None  # state the specified font colour for BODY ENGG1500 style and store in the variable CORRECT_FONT_COLOUR_BODY
    PARAGRAPH_STYLE = 'BODY ENGG1500'  # state the specified style name for paragraphs and store in the variable PARAGRAPH_STYLE
    for paragraph in WordFile.paragraphs:
        if PARAGRAPH_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # append font colours from each run into the set body_colour
                body_colour.add(run.font.color.rgb)
                # check if font colours are unacceptable, if so, store in the set body_wrong_colour
                if run.font.color.rgb is not CORRECT_FONT_COLOUR_BODY:
                    body_wrong_colour.add(run.font.color.rgb)
                    # store BODY ENGG1500 style text that contain unacceptable font colours in the sorted list body_wrong_colour_words
                    body_wrong_colour_words.append(run.text)

    # check if any elements in body_colour are not CORRECT_FONT_COLOUR_BODY
    if any(colour is not CORRECT_FONT_COLOUR_BODY for colour in body_colour) and len(body_colour) != 0:
        # print this if any elements in body_colour are not CORRECT_FONT_COLOUR_BODY and print body_wrong_colour and body_wrong_colour_words content
        st.error(f'''
        ❌ Text written in the BODY ENGG1500 style have incorrect font colour(s): {', '.join(map(str, body_wrong_colour))}  
        🡆 Incorrect font colours(s) found here: {' >> '.join(map(str, body_wrong_colour_words))}
        ''')
    elif len(body_colour) == 0:
        # print this if body_colour is empty, which means that BODY ENGG1500 style was not found
        st.info("ℹ️BODY ENGG1500 style font colour not found as this style was not used.")

    # TODO Font Colour code for BULLET LIST ENGG1500 style
    # check font colour for BULLET LIST ENGG1500 style (dot points) ********************
    bullet_colour = set()  # store all BULLET LIST ENGG1500 style font colours in the set bullet_colour
    bullet_wrong_colour = set()  # store unacceptable BULLET LIST ENGG1500 style font colours in the set bullet_wrong_colour
    bullet_wrong_colour_words = sorted(set())  # store BULLET LIST ENGG1500 style text that are in unacceptable font colour in the sorted list bullet_wrong_colour_words
    CORRECT_FONT_COLOUR_BULLET = None  # state the specified font colour for BULLET LIST ENGG1500 style and store in the variable CORRECT_FONT_COLOUR_BULLET
    DOT_POINT_STYLE = 'BULLET LIST ENGG1500'  # state the specified style name for dot points and store in the variable DOT_POINT_STYLE
    for paragraph in WordFile.paragraphs:
        if DOT_POINT_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add font colours from each run into the set bullet_colour
                bullet_colour.add(run.font.color.rgb)
                # check if font colours are unacceptable, if so, store in the set bullet_wrong_colour
                if run.font.color.rgb is not CORRECT_FONT_COLOUR_BULLET:
                    bullet_wrong_colour.add(run.font.color.rgb)
                    # append BULLET LIST ENGG1500 style text that contain unacceptable font colours in the sorted list bullet_wrong_colour_words
                    bullet_wrong_colour_words.append(run.text)

    # check if any elements in bullet_colour are not CORRECT_FONT_COLOUR_BULLET and if bullet_colour is not empty
    if any(colour is not CORRECT_FONT_COLOUR_BULLET for colour in bullet_colour) and len(bullet_colour) != 0:
        # print this if any elements in bullet_colour are not CORRECT_FONT_COLOUR_BULLET and print bullet_wrong_colour and bullet_wrong_colour_words content
        st.error(f'''
        ❌ Text written in the BULLET LIST ENGG1500 style have incorrect font colour(s): {', '.join(map(str, bullet_wrong_colour))}  
        🡆 Incorrect font colours(s) found here: {' >> '.join(map(str, bullet_wrong_colour_words))}
        ''')
    # check if bullet_colour is empty, if so this means that BULLET LIST ENGG1500 style was not found
    elif len(bullet_colour) == 0:
        # print this if bullet_colour is empty, which means that BULLET LIST ENGG1500 style was not found
        st.info("ℹ️BULLET LIST ENGG1500 style font colour not found as this style was not used.")

    # TODO Font Colour code for CAPTION ENGG1500 style
    # check font colour for CAPTION ENGG1500 style (figure and table captions) ********************
    cap_colour = set()  # store all CAPTION ENGG1500 style font colours in the set cap_colour
    cap_wrong_colour = set()  # store unacceptable CAPTION ENGG1500 style font colours in the set cap_wrong_colour
    cap_wrong_colour_words = sorted(set())  # store CAPTION ENGG1500 style text that are in unacceptable font colour in the sorted list cap_wrong_colour_words
    CORRECT_FONT_COLOUR_CAP = None  # state the specified font colour for CAPTION ENGG1500 style and store in the variable CORRECT_FONT_COLOUR_CAP
    CAPTION_STYLE = 'CAPTION ENGG1500'  # state the specified style name for figure and table captions and store in the variable CAPTION_STYLE
    for paragraph in WordFile.paragraphs:
        if CAPTION_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add font colours from each run into the set cap_colour
                cap_colour.add(run.font.color.rgb)
                # check if font colours are unacceptable, if so, store in the set cap_wrong_colour
                if run.font.color.rgb is not CORRECT_FONT_COLOUR_CAP:
                    cap_wrong_colour.add(run.font.color.rgb)
                    # append CAPTION ENGG1500 style text that contain unacceptable font colours in the sorted list cap_wrong_colour_words
                    cap_wrong_colour_words.append(run.text)

    # check if all elements in cap_colour are not CORRECT_FONT_COLOUR_CAP and if cap_colour is not empty
    if any(colour is not CORRECT_FONT_COLOUR_CAP for colour in cap_colour) and len(cap_colour) != 0:
        # print this if any elements in cap_colour are not CORRECT_FONT_COLOUR_CAP and print cap_wrong_colour and cap_wrong_colour_words content
        st.error(f'''
        ❌ Text written in the CAPTION ENGG1500 style have incorrect font colour(s): {', '.join(map(str, cap_wrong_colour))}  
        🡆 Incorrect font colours(s) found here: {' >> '.join(map(str, cap_wrong_colour_words))}
        ''')
    # check if cap_colour is empty, if so this means that CAPTION ENGG1500 style was not found
    elif len(cap_colour) == 0:
        # print this if cap_colour is empty, which means that CAPTION ENGG1500 style was not found
        st.info("ℹ️CAPTION ENGG1500 style font colour not found as this style was not used.")

    # TODO Font Colour code for TABLE ENGG1500 style
    # check font colour for TABLE ENGG1500 style (text in tables) ********************
    table_colour = set()  # store all TABLE ENGG1500 style font colours in the set table_colour
    table_wrong_colour = set()  # store unacceptable TABLE ENGG1500 style font colours in the set table_wrong_colour
    table_wrong_colour_words = sorted(set())  # store TABLE ENGG1500 style text that are in unacceptable font colour in the sorted list table_wrong_colour_words
    CORRECT_FONT_COLOUR_TABLE = None  # state the specified font colour for TABLE ENGG1500 style and store in the variable CORRECT_FONT_COLOUR_TABLE
    TABLE_TEXT_STYLE = 'TABLE ENGG1500'  # state the specified style name for text in tables and store in the variable TABLE_TEXT_STYLE
    for table in WordFile.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if TABLE_TEXT_STYLE == paragraph.style.name:
                        for run in paragraph.runs:
                            # add font colours from each run into the set table_colour
                            table_colour.add(run.font.color.rgb)
                            # check if fonts colours are unacceptable, if so, store in the set table_wrong_colour
                            if run.font.color.rgb is not CORRECT_FONT_COLOUR_TABLE:
                                table_wrong_colour.add(run.font.color.rgb)
                                # append TABLE ENGG1500 style text that contain unacceptable font colours in the sorted list table_wrong_colour_words
                                table_wrong_colour_words.append(run.text)

    # check if any elements in table_colour are not CORRECT_FONT_COLOUR_TABLE
    if any(colour is not CORRECT_FONT_COLOUR_TABLE for colour in table_colour) and len(table_colour) != 0:
        # print this if any elements in table_colour are not CORRECT_FONT_COLOUR_TABLE and print table_wrong_colour and table_wrong_colour_words content
        st.error(f'''
        ❌ Text written in the TABLE ENGG1500 style have incorrect font colour(s): {', '.join(map(str, table_wrong_colour))}  
        🡆 Incorrect font colours(s) found here: {' >> '.join(map(str, table_wrong_colour_words))}
        ''')
    # check if table_colour is empty, if so this means that TABLE ENGG1500 style was not found
    elif len(table_colour) == 0:
        # print this if table_colour is empty, which means that TABLE ENGG1500 style was not found
        st.info("ℹ️TABLE ENGG1500 style font colour not found as this style was not used.")


# Font Style program function
def font_style():
    # TODO start of Font Style code --------------------
    # add Font Style program banner
    st.subheader('**Font Style**')
    # access Word document file
    WordFile = Document(uploaded_file)

    # TODO Font Style code for TITLE ENGG1500 style
    # Check all TITLE ENGG1500 style text and determine if it is in bold and or italics ********************
    title_bold = set()  # store True and None for instances where TITLE ENGG1500 style is and not bold (None == off, True == on)
    title_italic = set()  # store True and None for instances where TITLE ENGG1500 style is and not italicised (None == off, True == on)
    title_not_bold_words = sorted(set())  # store TITLE ENGG1500 style text that are not in bold in the sorted list title_not_bold_words
    title_italic_words = sorted(set())  # store TITLE ENGG1500 style text that are italicised in the sorted list title_italic_words
    title_bold_italic_words = sorted(set())  # store TITLE ENGG1500 style text that are bold and italicised in the sorted list title_bold_italic_words
    TITLE_TEXT_STYLE = 'TITLE ENGG1500'  # state the specified style name for title page text and store in the variable TITLE_TEXT_STYLE
    for paragraph in WordFile.paragraphs:
        if TITLE_TEXT_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # append TITLE ENGG1500 style bold status from each run into the set title_bold
                title_bold.add(run.font.bold)
                # append TITLE ENGG1500 style italic status from each run into the set title_italic
                title_italic.add(run.font.italic)
                # check if TITLE ENGG1500 style is not bold
                if run.font.bold is False:
                    # append TITLE ENGG1500 style text that is not bold in the sorted list title_not_bold_words
                    title_not_bold_words.append(run.text)
                # check if TITLE ENGG1500 style is italicised
                if run.font.italic is True:
                    # append TITLE ENGG1500 style text that is italicised in the sorted list title_italic_words
                    title_italic_words.append(run.text)
                # check if TITLE ENGG1500 style is bold and italicised
                if run.font.bold is None and run.font.italic is True:
                    # append TITLE ENGG1500 style text that is bold and italicised in the sorted list title_bold_italic_words
                    title_bold_italic_words.append(run.text)

    # check if all elements in title_bold are None and if all elements in title_italic are None
    if all(t_style is False for t_style in title_bold) and all(t_style is None for t_style in title_italic) and len(title_bold) != 0 and len(title_italic) != 0:
        st.error("❌ All TITLE ENGG1500 style text are not bold.")
    # check if any of the elements of title_bold are None
    elif any(t_style is False for t_style in title_bold) and len(title_bold) != 0 and len(title_italic) != 0:
        st.error(f"❌ Non-bold TITLE ENGG1500 style text found: {' >> '.join(map(str, title_not_bold_words))}")

    # check if all elements of title_italic are True and if all elements of title_bold are None
    if all(t_style is True for t_style in title_italic) and all(t_style is False for t_style in title_bold) and len(title_bold) != 0 and len(title_italic) != 0:
        st.error("❌ All TITLE ENGG1500 style text are italicised.")
    # check if any of the elements of title_italic are True
    elif any(t_style is True for t_style in title_italic) and len(title_bold) != 0 and len(title_italic) != 0:
        st.error(f"❌ Italicised TITLE ENGG1500 style text found: {' >> '.join(map(str, title_italic_words))}")

    # check if all elements in title_bold and title_italic are True
    if all(t_style is None for t_style in title_bold) and all(t_style is True for t_style in title_italic) and len(title_bold) != 0 and len(title_italic) != 0:
        st.error("❌ All TITLE ENGG1500 style text are bold and italicised.")
    # check if any of the elements of title_bold and title_italic are True
    elif any(t_style is None for t_style in title_bold) and any(t_style is True for t_style in title_italic) and len(title_bold) != 0 and len(title_italic) != 0:
        st.error(f"❌ Bold and italicised TITLE ENGG1500 style text found: {' >> '.join(map(str, title_bold_italic_words))}")

    # check if both title_bold and title_italic are empty
    if len(title_bold) == 0 and len(title_italic) == 0:
        st.info("ℹ️TITLE ENGG1500 font style not found as it was not used.")

    # TODO Font Style code for HEADING 1 ENGG1500 style
    # Check all HEADING 1 ENGG1500 style text and determine if it is in bold and or italics ********************
    h1_bold = set()  # store True and None for instances where HEADING 1 ENGG1500 style is and not bold (None == off, True == on)
    h1_italic = set()  # store True and None for instances where HEADING 1 ENGG1500 style is and not italicised (None == off, True == on)
    h1_not_bold_words = sorted(set())  # store HEADING 1 ENGG1500 style text that are not in bold in the sorted list h1_not_bold_words
    h1_italic_words = sorted(set())  # store HEADING 1 ENGG1500 style text that are italicised in the sorted list h1_italic_words
    h1_bold_italic_words = sorted(set())  # store HEADING 1 ENGG1500 style text that are bold and italicised in the sorted list h1_bold_italic_words
    H1_STYLE = 'HEADING 1 ENGG1500'  # state the specified style name for main headings and store in the variable H1_STYLE
    for paragraph in WordFile.paragraphs:
        if H1_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # append HEADING 1 ENGG1500 style bold status from each run into the set h1_bold
                h1_bold.add(run.font.bold)
                # append HEADING 1 ENGG1500 style italic status from each run into the set h1_italic
                h1_italic.add(run.font.italic)
                # check if HEADING 1 ENGG1500 style is not bold
                if run.font.bold is False:
                    # append HEADING 1 ENGG1500 style text that is not bold in the sorted list h1_not_bold_words
                    h1_not_bold_words.append(run.text)
                # check if HEADING 1 ENGG1500 style is italicised
                if run.font.italic is True:
                    # append HEADING 1 ENGG1500 style text that is italicised in the sorted list h1_italic_words
                    h1_italic_words.append(run.text)
                # check if HEADING 1 ENGG1500 style is bold and italicised
                if run.font.bold is None and run.font.italic is True:
                    # append HEADING 1 ENGG1500 style text that is bold and italicised in the sorted list h1_bold_italic_words
                    h1_bold_italic_words.append(run.text)

    # check if all elements in h1_bold are None and if all elements in h1_italic are None
    if all(h1_style is False for h1_style in h1_bold) and all(h1_style is None for h1_style in h1_italic) and len(h1_bold) != 0 and len(h1_italic) != 0:
        st.error("❌ All HEADING 1 ENGG1500 style text are not bold.")
    # check if any of the elements of h1_bold are None
    elif any(h1_style is False for h1_style in h1_bold) and len(h1_bold) != 0 and len(h1_italic) != 0:
        st.error(f"❌ Non-bold HEADING 1 ENGG1500 style text found: {' >> '.join(map(str, h1_not_bold_words))}")

    # check if all elements of h1_italic are True and if all elements of h1_bold are None
    if all(h1_style is True for h1_style in h1_italic) and all(h1_style is False for h1_style in h1_bold) and len(h1_bold) != 0 and len(h1_italic) != 0:
        st.error("❌ All HEADING 1 ENGG1500 style text are italicised.")
    # check if any of the elements of h1_italic are True
    elif any(h1_style is True for h1_style in h1_italic) and len(h1_bold) != 0 and len(h1_italic) != 0:
        st.error(f"❌ Italicised HEADING 1 ENGG1500 style text found: {' >> '.join(map(str, h1_italic_words))}")

    # check if all elements in h1_bold and h1_italic are True
    if all(h1_style is None for h1_style in h1_bold) and all(h1_style is True for h1_style in h1_italic) and len(h1_bold) != 0 and len(h1_italic) != 0:
        st.error("❌ All HEADING 1 ENGG1500 style text are bold and italicised.")
    # check if any of the elements of h1_bold and h1_italic are True
    elif any(h1_style is None for h1_style in h1_bold) and any(h1_style is True for h1_style in h1_italic) and len(h1_bold) != 0 and len(h1_italic) != 0:
        st.error(f"❌ Bold and italicised HEADING 1 ENGG1500 style text found: {' >> '.join(map(str, h1_bold_italic_words))}")

    # check if both h1_bold and h1_italic are empty
    if len(h1_bold) == 0 and len(h1_italic) == 0:
        st.info("ℹ️HEADING 1 ENGG1500 font style not found as it was not used.")

    # TODO Font Style code for HEADING 2 ENGG1500 style
    # Check all Heading 2 HEADING 2 ENGG1500 style text and determine if it is in bold and or italics ********************
    h2_bold = set()  # store True and None for instances where HEADING 2 ENGG1500 is and not bold (None == off, True == on)
    h2_italic = set()  # store True and None for instances where HEADING 2 ENGG1500 is and not italicised (None == off, True == on)
    h2_not_italic_words = sorted(set())  # store HEADING 2 ENGG1500 style text that are not italicised in the sorted list h2_not_italic_words
    h2_bold_words = sorted(set())  # store HEADING 2 ENGG1500 style text that are bold in the sorted list h2_bold_words
    h2_bold_italic_words = sorted(set())  # store HEADING 2 ENGG1500 style text that are bold and italicised in the sorted list h2_bold_italic_words
    H2_STYLE = 'HEADING 2 ENGG1500'  # state the specified style name for sub headings and store in the variable H2_STYLE
    for paragraph in WordFile.paragraphs:
        if H2_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add HEADING 2 ENGG1500 style bold status from each run into the set h2_bold
                h2_bold.add(run.font.bold)
                # add HEADING 2 ENGG1500 style italic status from each run into the set h2_italic
                h2_italic.add(run.font.italic)
                # check if HEADING 2 ENGG1500 style is not italicised
                if run.font.italic is False:
                    # append HEADING 2 ENGG1500 style text that is not italicised in the sorted list h2_not_italic_words
                    h2_not_italic_words.append(run.text)
                # check if HEADING 2 ENGG1500 style is bold
                if run.font.bold is True:
                    # append HEADING 2 ENGG1500 style text that is bold in the sorted list h2_bold_words
                    h2_bold_words.append(run.text)
                # check if HEADING 2 ENGG1500 style is bold and italicised
                if run.font.bold is True and run.font.italic is None:
                    # append HEADING 2 ENGG1500 style text that is bold and italicised in the sorted list h2_bold_italic_words
                    h2_bold_italic_words.append(run.text)

    # check if all elements in h2_italic are None and if all elements in h2_bold are None
    if all(h2_style is False for h2_style in h2_italic) and all(h2_style is None for h2_style in h2_bold) and len(h2_bold) != 0 and len(h2_italic) != 0:
        st.error("❌ All HEADING 2 ENGG1500 style text are not italicised.")
    # check if any elements of h2_italic are None
    elif any(h2_style is False for h2_style in h2_italic) and len(h2_bold) != 0 and len(h2_italic) != 0:
        st.error(f"❌ Non-italicised HEADING 2 ENGG1500 style text found: {' >> '.join(map(str, h2_not_italic_words))}")

    # check if all elements in h2_bold are True and if all elements in h2_italic are None
    if all(h2_style is True for h2_style in h2_bold) and all(h2_style is False for h2_style in h2_italic) and len(h2_bold) != 0 and len(h2_italic) != 0:
        st.error("❌ All HEADING 2 ENGG1500 style text are bold.")
    # check if any of the elements of h2_bold are True
    elif any(h2_style is True for h2_style in h2_bold) and len(h2_bold) != 0 and len(h2_italic) != 0:
        st.error(f"❌ Bold HEADING 2 ENGG1500 style text found: {' >> '.join(map(str, h2_bold_words))}")

    # check if all elements in h2_bold and h2_italic are True
    if all(h2_style is True for h2_style in h2_bold) and all(h2_style is None for h2_style in h2_italic) and len(h2_bold) != 0 and len(h2_italic) != 0:
        st.error("❌ All HEADING 2 ENGG1500 style text are bold and italicised.")
    # check if any of the elements of h2_bold and h2_italic are True
    elif any(h2_style is True for h2_style in h2_bold) and any(h2_style is None for h2_style in h2_italic) and len(h2_bold) != 0 and len(h2_italic) != 0:
        st.error(f"❌ Bold and italicised HEADING 2 ENGG1500 style text found: {' >> '.join(map(str, h2_bold_italic_words))}")

    # check if both h2_bold and h2_italic are empty
    if len(h2_bold) == 0 and len(h2_italic) == 0:
        st.info("ℹ️HEADING 2 ENGG1500 font style not found as it was not used.")

    # TODO Font Style code for HEADING 3 ENGG1500 style
    # Check all HEADING 3 ENGG1500 style text and determine if it is in bold and or italics ********************
    h3_bold = set()  # store True and None for instances where HEADING 3 ENGG1500 style is and not bold (None == off, True == on)
    h3_italic = set()  # store True and None for instances where HEADING 3 ENGG1500 style is and not italicised (None == off, True == on)
    h3_not_italic_words = sorted(set())  # store HEADING 3 ENGG1500 style text that are not italicised in the sorted list h3_not_italic_words
    h3_bold_words = sorted(set())  # store HEADING 3 ENGG1500 style text that are bold in the sorted list h3_bold_words
    h3_bold_italic_words = sorted(set())  # store HEADING 3 ENGG1500 style text that are bold and italicised in the sorted list h3_bold_italic_words
    H3_STYLE = 'HEADING 3 ENGG1500'  # state the specified style name for sub headings and store in the variable H3_STYLE
    for paragraph in WordFile.paragraphs:
        if H3_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add HEADING 3 ENGG1500 style bold status from each run into the set h3_bold
                h3_bold.add(run.font.bold)
                # add HEADING 3 ENGG1500 style italic status from each run into the set h3_italic
                h3_italic.add(run.font.italic)
                # check if HEADING 3 ENGG1500 style is not italicised
                if run.font.italic is False:
                    # append HEADING 3 ENGG1500 style text that is not italicised in the sorted list h3_not_italic_words
                    h3_not_italic_words.append(run.text)
                # check if HEADING 3 ENGG1500 style is bold
                if run.font.bold is True:
                    # append HEADING 3 ENGG1500 style text that is bold in the sorted list h3_bold_words
                    h3_bold_words.append(run.text)
                # check if HEADING 3 ENGG1500 style is bold and italicised
                if run.font.bold is True and run.font.italic is None:
                    # append HEADING 3 ENGG1500 style text that is bold and italicised in the sorted list h3_bold_italic_words
                    h3_bold_italic_words.append(run.text)

    # check if all elements in h3_italic are None and if all elements in h3_bold are None
    if all(h3_style is False for h3_style in h3_italic) and all(h3_style is None for h3_style in h3_bold) and len(h3_bold) != 0 and len(h3_italic) != 0:
        st.error("❌ All HEADING 3 ENGG1500 style text are not italicised.")
    # check if any elements of h3_italic are None
    elif any(h3_style is False for h3_style in h3_italic) and len(h3_bold) != 0 and len(h3_italic) != 0:
        st.error(f"❌ Non-italicised HEADING 3 ENGG1500 style text found: {' >> '.join(map(str, h3_not_italic_words))}")

    # check if all elements in h3_bold are True and if all elements in h3_italic are None
    if all(h3_style is True for h3_style in h3_bold) and all(h3_style is False for h3_style in h3_italic) and len(h3_bold) != 0 and len(h3_italic) != 0:
        st.error("❌ All HEADING 3 ENGG1500 style text are bold.")
    # check if any of the elements of h3_bold are True
    elif any(h3_style is True for h3_style in h3_bold) and len(h3_bold) != 0 and len(h3_italic) != 0:
        st.error(f"❌ Bold HEADING 3 ENGG1500 style text found: {' >> '.join(map(str, h3_bold_words))}")

    # check if all elements in h3_bold and h3_italic are True
    if all(h3_style is True for h3_style in h3_bold) and all(h3_style is None for h3_style in h3_italic) and len(h3_bold) != 0 and len(h3_italic) != 0:
        st.error("❌ All HEADING 3 ENGG1500 style text are bold and italicised.")
    # check if any of the elements of h3_bold and h3_italic are True
    elif any(h3_style is True for h3_style in h3_bold) and any(h3_style is None for h3_style in h3_italic) and len(h3_bold) != 0 and len(h3_italic) != 0:
        st.error(f"❌ Bold and italicised HEADING 3 ENGG1500 style text found: {' >> '.join(map(str, h3_bold_italic_words))}")

    # check if both h3_bold and h3_italic are empty
    if len(h3_bold) == 0 and len(h3_italic) == 0:
        st.info("ℹ️HEADING 3 ENGG1500 font style not found as it was not used.")

    # TODO Font Style code for HEADING 4 ENGG1500 style
    # Check all HEADING 4 ENGG1500 style text and determine if it is in bold and or italics ********************
    h4_bold = set()  # store True and None for instances where HEADING 4 ENGG1500 style is and not bold (None == off, True == on)
    h4_italic = set()  # store None and False for instances where HEADING 4 ENGG1500 style is and not italicised (False == off, None == on)
    h4_not_italic_words = sorted(set())  # store HEADING 4 ENGG1500 style text that are not italicised in the sorted list h4_not_italic_words
    h4_bold_words = sorted(set())  # store HEADING 4 ENGG1500 style text that are bold in the sorted list h4_bold_words
    h4_bold_italic_words = sorted(set())  # store HEADING 4 ENGG1500 style text that are bold and italicised in the sorted list h4_bold_italic_words
    H4_STYLE = 'HEADING 4 ENGG1500'  # state the specified style name for sub headings and store in the variable H4_STYLE
    for paragraph in WordFile.paragraphs:
        if H4_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add HEADING 4 ENGG1500 style bold status from each run into the set h4_bold
                h4_bold.add(run.font.bold)
                # add HEADING 4 ENGG1500 style italic status from each run into the set h4_italic
                h4_italic.add(run.font.italic)
                # check if HEADING 4 ENGG1500 style is not italicised
                if run.font.italic is False:
                    # append HEADING 4 ENGG1500 style text that is not italicised in the sorted list h4_not_italic_words
                    h4_not_italic_words.append(run.text)
                # check if HEADING 4 ENGG1500 style is bold
                if run.font.bold is True:
                    # appendHEADING 4 ENGG1500 style text that is bold in the sorted list h4_bold_words
                    h4_bold_words.append(run.text)
                # check if HEADING 4 ENGG1500 style is bold and italicised
                if run.font.bold is True and run.font.italic is None:
                    # append HEADING 4 ENGG1500 style text that is bold and italicised in the sorted list h4_bold_italic_words
                    h4_bold_italic_words.append(run.text)

    # check if all elements in h4_italic are False and if all elements in h4_bold are None
    if all(h4_style is False for h4_style in h4_italic) and all(h4_style is None for h4_style in h4_bold) and len(h4_bold) != 0 and len(h4_italic) != 0:
        st.error("❌ All HEADING 4 ENGG1500 style text are not italicised.")
    # check if any elements of h4_italic are False
    elif any(h4_style is False for h4_style in h4_italic) and len(h4_bold) != 0 and len(h4_italic) != 0:
        st.error(f"❌ Non-italicised HEADING 4 ENGG1500 style text found: {' >> '.join(map(str, h4_not_italic_words))}")

    # check if all elements in h4_bold are True and if all elements in h4_italic are False
    if all(h4_style is True for h4_style in h4_bold) and all(h4_style is False for h4_style in h4_italic) and len(h4_bold) != 0 and len(h4_italic) != 0:
        st.error("❌ All HEADING 4 ENGG1500 style text are bold.")
    # check if any of the elements of h4_bold are True
    elif any(h4_style is True for h4_style in h4_bold) and len(h4_bold) != 0 and len(h4_italic) != 0:
        st.error(f"❌ Bold HEADING 4 ENGG1500 style text found: {' >> '.join(map(str, h4_bold_words))}")

    # check if all elements in h4_bold and h4_italic are True and None, respectively
    if all(h4_style is True for h4_style in h4_bold) and all(h4_style is None for h4_style in h4_italic) and len(h4_bold) != 0 and len(h4_italic) != 0:
        st.error("❌ All HEADING 4 ENGG1500 style text are bold and italicised.")
    # check if any of the elements in h4_bold and h4_italic are True and None, respectively
    elif any(h4_style is True for h4_style in h4_bold) and any(h4_style is None for h4_style in h4_italic) and len(h4_bold) != 0 and len(h4_italic) != 0:
        st.error(f"❌ Bold and italicised HEADING 4 ENGG1500 style text found: {' >> '.join(map(str, h4_bold_italic_words))}")

    # check if both h4_bold and h4_italic are empty
    if len(h4_bold) == 0 and len(h4_italic) == 0:
        st.info("ℹ️HEADING 4 ENGG1500 font style not found as it was not used.")


# paragraph alignment program function
def para_align():
    # TODO start of Paragraph Alignment code --------------------
    # add paragraph alignment program banner
    st.subheader("**Paragraph Alignment**")
    # access Word document file
    WordFile = Document(uploaded_file)

    # TODO Paragraph Alignment code for BODY ENGG1500 style
    # check the paragraph alignment of all text formatted with the BODY ENGG1500 style (paragraphs) ********************
    body_align = set()  # store all paragraph alignments in the set body_align
    body_wrong_align = set()  # store unacceptable paragraph alignments in the set body_wrong_align
    body_wrong_align_words = sorted(set())  # store text from unacceptable paragraph alignments in the sorted list body_wrong_align
    CORRECT_ALIGN_BODY = None  # state the specified alignment for BODY ENGG1500 style and store in the variable CORRECT_ALIGN_BODY
    PARAGRAPH_STYLE = 'BODY ENGG1500'  # state the specified style name for paragraphs and store in the variable PARAGRAPH_STYLE
    for paragraph in WordFile.paragraphs:
        if PARAGRAPH_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add paragraph alignment from each run into the set body_align
                body_align.add(paragraph.alignment)
                # check if paragraph alignment is unacceptable, if so, store in the set body_wrong_align
                if paragraph.alignment is not CORRECT_ALIGN_BODY:
                    body_wrong_align.add(paragraph.alignment)
                    # append BODY ENGG1500 style text that contain unacceptable paragraph alignment in the sorted list body_wrong_align_words
                    if paragraph.text not in body_wrong_align_words:
                        body_wrong_align_words.append(paragraph.text)

    # check if any elements in body_align are not CORRECT_ALIGN_BODY
    if any(alignment is not CORRECT_ALIGN_BODY for alignment in body_align) and len(body_align) != 0:
        # print this if any elements in body_align are not CORRECT_ALIGN_BODY and print body_wrong_align and body_wrong_align_words contents
        st.error(f'''
        ❌ Text written in the BODY ENGG1500 style have incorrect paragraph alignment: {', '.join(map(str, body_wrong_align))}  
        🡆 Incorrect paragraph alignment found here: {' >> '.join(map(str, body_wrong_align_words))}
        ''')
    # check if the set body_align is empty
    elif len(body_align) == 0:
        st.info("ℹ️BODY ENGG1500 style paragraph alignment not found as this style was not used.")

    # TODO Paragraph Alignment code for BULLET LIST ENGG1500 style
    # check the paragraph alignment of all text formatted with the BULLET LIST ENGG1500 style (dot points) ********************
    bullet_align = set()  # store all paragraph alignments in the set bullet_align
    bullet_wrong_align = set()  # store unacceptable paragraph alignments in the set bullet_wrong_align
    bullet_wrong_align_words = sorted(set())  # store text from unacceptable BULLET LIST ENGG1500 style alignments in the sorted list bullet_wrong_align
    CORRECT_ALIGN_BULLET = None  # state the specified alignment for BULLET LIST ENGG1500 style and store in the variable CORRECT_ALIGN_BULLET
    DOT_POINT_STYLE = 'BULLET LIST ENGG1500'  # state the specified style name for dot points and store in the variable DOT_POINT_STYLE
    for paragraph in WordFile.paragraphs:
        if DOT_POINT_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add paragraph alignment from each run into the set bullet_align
                bullet_align.add(paragraph.alignment)
                # check if paragraph alignment is unacceptable, if so, store in the set bullet_wrong_align
                if paragraph.alignment is not CORRECT_ALIGN_BULLET:
                    bullet_wrong_align.add(paragraph.alignment)
                    # append BULLET LIST ENGG1500 style text that contain unacceptable paragraph alignment in the sorted list bullet_wrong_align_words
                    if paragraph.text not in bullet_wrong_align_words:
                        bullet_wrong_align_words.append(paragraph.text)

    # check if any elements in bullet_align are not CORRECT_ALIGN_BULLET and if the list bullet_align is not empty
    if any(alignment is not CORRECT_ALIGN_BULLET for alignment in bullet_align) and len(bullet_align) != 0:
        # print this if any elements in bullet_align are not CORRECT_ALIGN_BULLET and print bullet_wrong_align and bullet_wrong_align_words contents
        st.error(f'''
        ❌ Text written in the BULLET LIST ENGG1500 style have incorrect paragraph alignment: {', '.join(map(str, bullet_wrong_align))}  
        🡆 Incorrect paragraph alignment found here: {' >> '.join(map(str, bullet_wrong_align_words))}
        ''')
    # check if the set bullet_align is empty
    elif len(bullet_align) == 0:
        st.info("ℹ️BULLET LIST ENGG1500 style paragraph alignment not found as this style was not used.")

    # TODO Paragraph Alignment code for TABLE ENGG1500 style
    # check the paragraph alignment of all text formatted with the TABLE ENGG1500 style (text in tables) ********************
    table_align = set()  # store all TABLE ENGG1500 style paragraph alignments in the set table_align
    table_wrong_align = set()  # store unacceptable TABLE ENGG1500 style paragraph alignments in the set table_wrong_align
    table_wrong_align_words = sorted(set())  # store TABLE ENGG1500 style text with unacceptable paragraph alignments in the sorted list table_wrong_align_words
    CORRECT_ALIGN_TABLE = None  # state the specified alignment for TABLE ENGG1500 style and store in the variable CORRECT_ALIGN_TABLE
    TABLE_TEXT_STYLE = 'TABLE ENGG1500'  # state the specified style name for text in tables and store in the variable TABLE_TEXT_STYLE
    for table in WordFile.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if TABLE_TEXT_STYLE == paragraph.style.name:
                        for run in paragraph.runs:
                            # add paragraph alignment from each run into the set table_align
                            table_align.add(paragraph.alignment)
                            # check if paragraph alignment is unacceptable, if so, store in the set table_wrong_align
                            if paragraph.alignment is not CORRECT_ALIGN_TABLE:
                                table_wrong_align.add(paragraph.alignment)
                                # append TABLE ENGG1500 style text that contain unacceptable paragraph alignment in the sorted list table_wrong_align_words
                                if paragraph.text not in table_wrong_align_words:
                                    table_wrong_align_words.append(paragraph.text)

    # check if any elements in table_align are not CORRECT_ALIGN_TABLE
    if any(alignment is not CORRECT_ALIGN_TABLE for alignment in table_align) and len(table_align) != 0:
        # print this if any elements in table_align are not CORRECT_ALIGN_TABLE and print table_wrong_align and table_wrong_align_words contents
        st.error(f'''
        ❌ Text written in the TABLE ENGG1500 style have incorrect paragraph alignment: {', '.join(map(str, table_wrong_align))}  
        🡆 Incorrect paragraph alignment found here: {' >> '.join(map(str, table_wrong_align_words))}
        ''')
    # check if the set table_align is empty
    elif len(table_align) == 0:
        st.info("ℹ️TABLE ENGG1500 style paragraph alignment not found as this style was not used.")


# spacing program function
def spacing():
    # TODO start of Spacing code --------------------
    # add spacing program banner
    st.subheader("**Spacing**")
    # access Word document file
    WordFile = Document(uploaded_file)

    # TODO spacing before code for BODY ENGG1500 style
    # Check spacing before paragraph for all text formatted in the BODY ENGG1500 style (paragraphs) ********************
    body_para_b = set()  # store all spacing before paragraph values for BODY ENGG1500 style text in the set body_para_b
    body_wrong_para_b = set()  # store unacceptable spacing before paragraph values for BODY ENGG1500 style text in the set body_wrong_para_b
    body_wrong_para_b_words = sorted(set())  # store BODY ENGG1500 style text that contain unacceptable spacing before paragraph values in the sorted set body_wrong_para_b_words
    CORRECT_SPACE_BEFORE_BODY = None  # state the specified spacing before for BODY ENGG1500 style and store in the variable CORRECT_SPACE_BEFORE_BODY
    PARAGRAPH_STYLE = 'BODY ENGG1500'  # state the specified style name for paragraphs and store in the variable PARAGRAPH_STYLE
    for paragraph in WordFile.paragraphs:
        if PARAGRAPH_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add BODY ENGG1500 style text spacing before paragraph values in the set body_para_b
                body_para_b.add(paragraph.paragraph_format.space_before)
                # check if spacing before paragraph is unacceptable, if so, store in the set body_wrong_para_b
                if paragraph.paragraph_format.space_before is not CORRECT_SPACE_BEFORE_BODY:
                    body_wrong_para_b.add(paragraph.paragraph_format.space_before/12700)
                    # append BODY ENGG1500 style text that contain unacceptable spacing before paragraph values in the sorted list body_wrong_para_b_words
                    body_wrong_para_b_words.append(run.text)

    # check if any elements in body_para_b are not CORRECT_SPACE_BEFORE_BODY
    if any(space is not CORRECT_SPACE_BEFORE_BODY for space in body_para_b) and len(body_para_b) != 0:
        # print this if any elements in body_para_b are not CORRECT_SPACE_BEFORE_BODY and print body_wrong_para_b content
        st.error(f'''
        ❌ Text written in the BODY ENGG1500 style have incorrect spacing before paragraph: {', '.join(map(str, body_wrong_para_b))}  
        🡆 Incorrect spacing before paragraph found here: {' >> '.join(map(str, body_wrong_para_b_words))}
        ''')
    # check if body_para_b is empty, if so this means that BODY ENGG1500 style was not found
    elif len(body_para_b) == 0:
        # print this if body_para_b is empty, since BODY ENGG1500 style was not found
        st.info("ℹ️BODY ENGG1500 style spacing before paragraph not found as this style was not used.")

    # TODO line spacing code for BODY ENGG1500 style
    # Check line spacing for all text formatted in the BODY ENGG1500 style (paragraphs) ***************************
    body_para_l = set()  # store all BODY ENGG1500 style text line spacing values in the set body_para_l
    body_wrong_para_l = set()  # store unacceptable BODY ENGG1500 style text line spacing values in the set body_wrong_para_l
    body_wrong_para_l_words = sorted(set())  # store BODY ENGG1500 style text with unacceptable line spacing in sorted list body_wrong_para_l_words
    CORRECT_LINE_SPACE_BODY = None  # state the specified line spacing for BODY ENGG1500 style and store in the variable CORRECT_LINE_SPACE_BODY
    PARAGRAPH_STYLE = 'BODY ENGG1500'  # state the specified style name for paragraphs and store in the variable PARAGRAPH_STYLE
    for paragraph in WordFile.paragraphs:
        if PARAGRAPH_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add BODY ENGG1500 style text line spacing values in the set body_para_l
                body_para_l.add(paragraph.paragraph_format.line_spacing)
                # check if BODY ENGG1500 style text line spacing is unacceptable, if so, store in the set body_wrong_para_l
                if paragraph.paragraph_format.line_spacing is not CORRECT_LINE_SPACE_BODY:
                    body_wrong_para_l.add(paragraph.paragraph_format.line_spacing)
                    # append BODY ENGG1500 style text that contain unacceptable line spacing in the sorted list body_wrong_para_l_words
                    body_wrong_para_l_words.append(run.text)

    # check if any elements in body_para_l are not CORRECT_LINE_SPACE_BODY
    if any(space is not CORRECT_LINE_SPACE_BODY for space in body_para_l) and len(body_para_l) != 0:
        # print this if any elements in body_para_l are not CORRECT_LINE_SPACE_BODY and print body_wrong_para_l and body_wrong_para_l_words content
        st.error(f'''
        ❌ Text written in the BODY ENGG1500 style have incorrect line spacing: {', '.join(map(str, body_wrong_para_l))}  
        🡆 Incorrect line spacing found here: {' >> '.join(map(str, body_wrong_para_l_words))}
        ''')
    # check if body_para_l is empty, if so this means that BODY ENGG1500 style was not found
    elif len(body_para_l) == 0:
        # print this if body_para_l is empty, since BODY ENGG1500 style was not found
        st.info("ℹ️BODY ENGG1500 style line spacing not found as this style was not used.")

    # TODO spacing after code for BODY ENGG1500 style
    # Check spacing after paragraph for all text formatted in the BODY ENGG1500 style (paragraphs) ********************
    body_para_a = set()  # store all spacing after paragraph values for BODY ENGG1500 style text in the set body_para_a
    body_wrong_para_a = set()  # store unacceptable spacing after paragraph values for BODY ENGG1500 style text in the set body_wrong_para_a
    body_wrong_para_a_words = sorted(set())  # store BODY ENGG1500 style text that contain unacceptable spacing after paragraph values in the sorted set body_wrong_para_a_words
    CORRECT_SPACE_AFTER_BODY = None  # state the specified spacing after for BODY ENGG1500 style and store in the variable CORRECT_SPACE_AFTER_BODY
    PARAGRAPH_STYLE = 'BODY ENGG1500'  # state the specified style name for paragraphs and store in the variable PARAGRAPH_STYLE
    for paragraph in WordFile.paragraphs:
        if PARAGRAPH_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add BODY ENGG1500 style text spacing after paragraph values in the set body_para_a
                body_para_a.add(paragraph.paragraph_format.space_after)
                # check if spacing after paragraph is unacceptable, if so, store in the set body_wrong_para_a
                if paragraph.paragraph_format.space_after is not CORRECT_SPACE_AFTER_BODY:
                    body_wrong_para_a.add(paragraph.paragraph_format.space_after/12700)
                    # append BODY ENGG1500 style text that contain unacceptable spacing after paragraph values in the sorted list body_wrong_para_a_words
                    body_wrong_para_a_words.append(run.text)

    # check if any elements in body_para_a are not CORRECT_SPACE_AFTER_BODY
    if any(space is not CORRECT_SPACE_AFTER_BODY for space in body_para_a) and len(body_para_a) != 0:
        # print this if any elements in body_para_a are not CORRECT_SPACE_AFTER_BODY and print body_wrong_para_a and body_wrong_para_a_words content
        st.error(f'''
        ❌ Text written in the BODY ENGG1500 style have incorrect spacing after paragraph: {', '.join(map(str, body_wrong_para_a))}  
        🡆 Incorrect spacing after paragraph found here: {' >> '.join(map(str, body_wrong_para_a_words))}
        ''')
    # check if body_para_a is empty, if so this means that BODY ENGG1500 style was not found
    elif len(body_para_a) == 0:
        # print this if body_para_a is empty, since BODY ENGG1500 style was not found
        st.info("ℹ️BODY ENGG1500 style spacing after paragraph not found as this style was not used.")

    # TODO spacing before code for BULLET LIST ENGG1500 style
    # Check spacing before paragraph for all text formatted in the BULLET LIST ENGG1500 style (dot points) ********************
    bullet_para_b = set()  # store all spacing before paragraph values for BULLET LIST ENGG1500 style text in the set bullet_para_b
    bullet_wrong_para_b = set()  # store unacceptable spacing before paragraph values for BULLET LIST ENGG1500 style text in the set bullet_wrong_para_b
    bullet_wrong_para_b_words = sorted(set())  # store BULLET LIST ENGG1500 style text that contain unacceptable spacing before paragraph values in the sorted list bullet_wrong_para_b_words
    CORRECT_SPACE_BEFORE_BULLET = None  # state the specified spacing before for BULLET LIST ENGG1500 style and store in the variable CORRECT_SPACE_BEFORE_BULLET
    DOT_POINT_STYLE = 'BULLET LIST ENGG1500'  # state the specified style name for dot points and store in the variable DOT_POINT_STYLE
    for paragraph in WordFile.paragraphs:
        if DOT_POINT_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add BULLET LIST ENGG1500 style text spacing before paragraph values in the set bullet_para_b
                bullet_para_b.add(paragraph.paragraph_format.space_before)
                # check if spacing before paragraph is unacceptable, if so, store in the set bullet_wrong_para_b
                if paragraph.paragraph_format.space_before is not CORRECT_SPACE_BEFORE_BULLET:
                    bullet_wrong_para_b.add(paragraph.paragraph_format.space_before/12700)
                    # append BULLET LIST ENGG1500 style text that contain unacceptable spacing before paragraph values in the sorted list bullet_wrong_para_b_words
                    bullet_wrong_para_b_words.append(run.text)

    # check if any elements in bullet_para_b are not CORRECT_SPACE_BEFORE_BULLET
    if any(space is not CORRECT_SPACE_BEFORE_BULLET for space in bullet_para_b) and len(bullet_para_b) != 0:
        # print this if any elements in bullet_para_b are not CORRECT_SPACE_BEFORE_BULLET and print bullet_wrong_para_b and bullet_wrong_para_b_words content
        st.error(f'''
        ❌ Text written in the BULLET LIST ENGG1500 style have incorrect spacing before paragraph: {', '.join(map(str, bullet_wrong_para_b))}  
        🡆 Incorrect spacing before paragraph found here: {' >> '.join(map(str, bullet_wrong_para_b_words))}
        ''')
    # check if bullet_para_b is empty, if so this means that BULLET LIST ENGG1500 style was not found
    elif len(bullet_para_b) == 0:
        # print this if bullet_para_b is empty, since BULLET LIST ENGG1500 style was not found
        st.info("ℹ️BULLET LIST ENGG1500 style spacing before paragraph not found as this style was not used.")

    # TODO line spacing code for BULLET LIST ENGG1500 style
    # Check line spacing for all text formatted in the BULLET LIST ENGG1500 style (bullet list) ***************************
    bullet_para_l = set()  # store all BULLET LIST ENGG1500 style text line spacing values in the set bullet_para_l
    bullet_wrong_para_l = set()  # store unacceptable BULLET LIST ENGG1500 style text line spacing values in the set bullet_wrong_para_l
    bullet_wrong_para_l_words = sorted(set())  # store BULLET LIST ENGG1500 style text with unacceptable line spacing in the sorted list bullet_wrong_para_l_words
    CORRECT_LINE_SPACE_BULLET = None  # state the specified line spacing for BULLET LIST ENGG1500 style and store in the variable CORRECT_LINE_SPACE_BULLET
    DOT_POINT_STYLE = 'BULLET LIST ENGG1500'  # state the specified style name for dot points and store in the variable DOT_POINT_STYLE
    for paragraph in WordFile.paragraphs:
        if DOT_POINT_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add BULLET LIST ENGG1500 style text line spacing values in the set bullet_para_l
                bullet_para_l.add(paragraph.paragraph_format.line_spacing)
                # check if line spacing are unacceptable, if so, store in the set bullet_wrong_para_l
                if paragraph.paragraph_format.line_spacing is not CORRECT_LINE_SPACE_BULLET:
                    bullet_wrong_para_l.add(paragraph.paragraph_format.line_spacing)
                    # append BULLET LIST ENGG1500 style text that contain unacceptable line spacing in the sorted list bullet_wrong_para_l_words
                    bullet_wrong_para_l_words.append(run.text)

    # check if any elements in bullet_para_l are not CORRECT_LINE_SPACE_BULLET
    if any(space is not CORRECT_LINE_SPACE_BULLET for space in bullet_para_l) and len(bullet_para_l) != 0:
        # print this if any elements in bullet_para_l are not CORRECT_LINE_SPACE_BULLET and print bullet_wrong_para_l and bullet_wrong_para_l_words content
        st.error(f'''
        ❌ Text written in the BULLET LIST ENGG1500 style have incorrect line spacing: {', '.join(map(str, bullet_wrong_para_l))}  
        🡆 Incorrect line spacing found here: {' >> '.join(map(str, bullet_wrong_para_l_words))}
        ''')
    # check if bullet_para_l is empty, if so this means that BULLET LIST ENGG1500 style was not found
    elif len(bullet_para_l) == 0:
        # print this if bullet_para_l is empty, since BULLET LIST ENGG1500 style was not found
        st.info("ℹ️BULLET LIST ENGG1500 style line spacing not found as this style was not used.")

    # TODO spacing after code for BULLET LIST ENGG1500 style
    # Check spacing after paragraph for all text formatted in the BULLET LIST ENGG1500 style (dot points) ********************
    bullet_para_a = set()  # store all spacing after paragraph values for BULLET LIST ENGG1500 style text in the set bullet_para_a
    bullet_wrong_para_a = set()  # store unacceptable spacing after values for BULLET LIST ENGG1500 style text in the set bullet_wrong_para_a
    bullet_wrong_para_a_words = sorted(set())  # store BULLET LIST ENGG1500 style text that contain unacceptable spacing after paragraph values in the sorted list bullet_wrong_para_a_words
    CORRECT_SPACE_AFTER_BULLET = None  # state the specified spacing after for BULLET LIST ENGG1500 style and store in the variable CORRECT_SPACE_AFTER_BULLET
    DOT_POINT_STYLE = 'BULLET LIST ENGG1500'  # state the specified style name for dot points and store in the variable DOT_POINT_STYLE
    for paragraph in WordFile.paragraphs:
        if DOT_POINT_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add BULLET LIST ENGG1500 style text spacing after paragraph values in the set bullet_para_a
                bullet_para_a.add(paragraph.paragraph_format.space_after)
                # check if spacing after paragraph is unacceptable, if so, store in the set bullet_wrong_para_a
                if paragraph.paragraph_format.space_after is not CORRECT_SPACE_AFTER_BULLET:
                    bullet_wrong_para_a.add(paragraph.paragraph_format.space_after/12700)
                    # append BULLET LIST ENGG1500 style text that contain unacceptable spacing after paragraph values in the sorted list bullet_wrong_para_a_words
                    bullet_wrong_para_a_words.append(run.text)

    # check if any elements in bullet_para_a are not CORRECT_SPACE_AFTER_BULLET
    if any(space is not CORRECT_SPACE_AFTER_BULLET for space in bullet_para_a) and len(bullet_para_a) != 0:
        # print this if any elements in bullet_para_a are not CORRECT_SPACE_AFTER_BULLET and print bullet_wrong_para_a and bullet_wrong_para_a_words content
        st.error(f'''
        ❌ Text written in the BULLET LIST ENGG1500 style have incorrect spacing after paragraph: {', '.join(map(str, bullet_wrong_para_a))}  
        🡆 Incorrect spacing after paragraph found here: {' >> '.join(map(str, bullet_wrong_para_a_words))}
        ''')
    # check if bullet_para_a is empty, if so this means that BULLET LIST ENGG1500 style was not found
    elif len(bullet_para_a) == 0:
        # print this if bullet_para_a is empty, since BULLET LIST ENGG1500 style was not found
        st.info("ℹ️BULLET LIST ENGG1500 style spacing after paragraph not found as this style was not used.")

    # TODO line spacing code for CAPTION ENGG1500 style
    # Check line spacing for all text formatted in the CAPTION ENGG1500 style (figure and table captions) ***************************
    cap_para_l = set()  # store all CAPTION ENGG1500 style text line spacing values in the set cap_para_l
    cap_wrong_para_l = set()  # store unacceptable CAPTION ENGG1500 style text line spacing values in the set cap_wrong_para_l
    cap_wrong_para_l_words = sorted(set())  # store CAPTION ENGG1500 style text with unacceptable line spacing in the sorted list cap_wrong_para_l_words
    CORRECT_LINE_SPACE_CAP = None  # state the specified line spacing for CAPTION ENGG1500 style and store in the variable CORRECT_LINE_SPACE_CAP
    CAPTION_STYLE = 'CAPTION ENGG1500'  # state the specified style name for figure and table captions and store in the variable CAPTION_STYLE
    for paragraph in WordFile.paragraphs:
        if CAPTION_STYLE == paragraph.style.name:
            for run in paragraph.runs:
                # add CAPTION ENGG1500 style text line spacing values in the set cap_para_l
                cap_para_l.add(paragraph.paragraph_format.line_spacing)
                # check if line spacing is unacceptable, if so, store in the set cap_wrong_para_l
                if paragraph.paragraph_format.line_spacing is not CORRECT_LINE_SPACE_CAP:
                    cap_wrong_para_l.add(paragraph.paragraph_format.line_spacing)
                    # append CAPTION ENGG1500 style text that contain unacceptable line spacing in the sorted list cap_wrong_para_l_words
                    cap_wrong_para_l_words.append(run.text)

    # check if any elements in cap_para_l are not CORRECT_LINE_SPACE_CAP
    if any(space is not CORRECT_LINE_SPACE_CAP for space in cap_para_l) and len(cap_para_l) != 0:
        # print this if any elements in cap_para_l are not CORRECT_LINE_SPACE_CAP and print cap_wrong_para_l and cap_wrong_para_l_words content
        st.error(f'''
        ❌ Text written in the CAPTION ENGG1500 style have incorrect line spacing: {', '.join(map(str, cap_wrong_para_l))}  
        🡆 Incorrect line spacing found here: {' >> '.join(map(str, cap_wrong_para_l_words))}
        ''')
    # check if cap_para_l is empty, if so this means that CAPTION ENGG1500 style was not found
    elif len(cap_para_l) == 0:
        # print this if cap_para_l is empty, since CAPTION ENGG1500 style was not found
        st.info("ℹ️CAPTION ENGG1500 style line spacing not found as this style was not used.")

    # TODO line spacing code for TABLE ENGG1500 style
    # Check line spacing for all text formatted in the TABLE ENGG1500 style (text in tables) ***************************
    table_para_l = set()  # store all TABLE ENGG1500 style text line spacing values in the set table_para_l
    table_wrong_para_l = set()  # store unacceptable TABLE ENGG1500 style text line spacing values in the set table_wrong_para_l
    table_wrong_para_l_words = sorted(set())  # store TABLE ENGG1500 style text with unacceptable line spacing in sorted list table_wrong_para_l_words
    CORRECT_LINE_SPACE_TABLE = None  # state the specified line spacing for TABLE ENGG1500 style and store in the variable CORRECT_LINE_SPACE_TABLE
    TABLE_TEXT_STYLE = 'TABLE ENGG1500'  # state the specified style name for text in tables and store in the variable TABLE_TEXT_STYLE
    for table in WordFile.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if TABLE_TEXT_STYLE == paragraph.style.name:
                        for run in paragraph.runs:
                            # add TABLE ENGG1500 style text line spacing values in the set table_para_l
                            table_para_l.add(paragraph.paragraph_format.line_spacing)
                            # check if TABLE ENGG1500 style text line spacing is unacceptable, if so, store in the set table_wrong_para_l
                            if paragraph.paragraph_format.line_spacing is not CORRECT_LINE_SPACE_TABLE:
                                table_wrong_para_l.add(paragraph.paragraph_format.line_spacing)
                                # append TABLE ENGG1500 style text that contain unacceptable line spacing in the sorted list table_wrong_para_l_words
                                table_wrong_para_l_words.append(run.text)

    # check if any elements in table_para_l are not CORRECT_LINE_SPACE_TABLE
    if any(space is not CORRECT_LINE_SPACE_TABLE for space in table_para_l) and len(table_para_l) != 0:
        # print this if any elements in table_para_l are not CORRECT_LINE_SPACE_TABLE and print table_wrong_para_l and table_wrong_para_l_words content
        st.error(f'''
        ❌ Text written in the TABLE ENGG1500 style have incorrect line spacing: {', '.join(map(str, table_wrong_para_l))}  
        🡆 Incorrect line spacing found here: {' >> '.join(map(str, table_wrong_para_l_words))}
        ''')
    # check if table_para_l is empty, if so this means that TABLE ENGG1500 style was not found
    elif len(table_para_l) == 0:
        # print this if table_para_l is empty, since TABLE ENGG1500 style was not found
        st.info("ℹ️TABLE ENGG1500 style line spacing not found as this style was not used.")


# reference counter program function
def ref_count():
    # TODO start of Reference Counter code --------------------
    # add reference counter program banner
    st.subheader("**Reference Counter**")
    # Access the Word document
    TextDoc_ieee = docx.process(uploaded_file)

    # Count the number of IEEE citations in the document ********************
    # Regex pattern to find IEEE style citations
    pattern_ieee = r'(\[\d+-\d+\]|\[\d+(,\d+)*\])'
    # Try to find matches (returned as an iterator of matches) for IEEE
    results_ieee = re.finditer(pattern_ieee, TextDoc_ieee)

    # Build a list with IEEE citations obtained by looping through matches
    # Each match has the first and last indices of the match, relative to the original string
    references_ieee = [TextDoc_ieee[match.start(): match.end()] for match in results_ieee]
    # Remove duplicate citations
    unique_citations = list(set(references_ieee))
    # store number of unique citations as an integer in cite_num_int
    cite_num_int = len(unique_citations)
    # store number of unique citations as a string in cite_num_str
    cite_num_str = str(cite_num_int)

    # Check if the length of the list unique_citations is less than 5 and if it is not empty
    if len(unique_citations) < 5 and len(unique_citations) != 0:
        st.error(f'''
        ❌ Less than five IEEE style references were found.  
        🡆 Number of IEEE style references: {cite_num_str}
        ''')

    # Check if the list unique_citations is empty, meaning no IEEE references were found
    elif len(unique_citations) == 0:
        st.error("❌ No IEEE style references were found.")


# header and footer program function
def header_footer():
    # TODO start of Header and Footer code --------------------
    # add header and footer program banner
    st.subheader("**Header and Footer**")
    # access Word document file
    WordFile = Document(uploaded_file)

    # Extract headers and footers from document *******************
    headers = []  # store page headers in the list headers
    # loop through all sections in the document
    for section in WordFile.sections:
        for paragraph in section.header.paragraphs:
            # only append text from headers if it is not an empty string
            if '' != paragraph.text:
                headers.append(paragraph.text)

    # check if length of the list headers is not equal to zero, which implies headers exist in the document
    if len(headers) != 0:
        st.write(f"The following headers were found: {' >> '.join(map(str, headers))}")
    # check if length of the list headers is equal to zero, which implies headers do not exist in the document
    elif len(headers) == 0:
        st.error("❌ No headers were found.")

    footers = []  # store page footers in the list footers
    # loop through all sections in the document
    for section in WordFile.sections:
        for paragraph in section.footer.paragraphs:
            # only append text from footers if it is not an empty string
            if '' != paragraph.text:
                footers.append(paragraph.text)

    # check if length of the list footers is not equal to zero, which implies footers exist in the document
    if len(footers) != 0:
        st.write(f"The following footers were found: {' >> '.join(map(str, footers))}")
    # check if length of the list footers is equal to zero, which implies footers do not exist in the document
    elif len(footers) == 0:
        st.error("❌ No footers were found.")


# page margins program function
def page_margin():
    # TODO start of Page Margins code --------------------
    # add page margins program banner
    st.subheader("**Page Margins**")
    # access Word document file
    WordFile = Document(uploaded_file)

    # scan the document and check if page margins are valid ********************
    page_margins = []  # store page margins in the list page_margins
    # loop through all sections in the document
    for section in WordFile.sections:
        # append page margins in the list page_margins
        page_margins.append((section.top_margin, section.bottom_margin, section.left_margin, section.right_margin))

    # execute function to check if page margins correspond to those of Narrow margins
    # takes a value (margins) and returns true if it corresponds to those of Narrow margins
    def Narrow_present(margins):
        TOP_MARGIN = 457200  # top margin value for Narrow margins as interpreted by python-docx
        BOTTOM_MARGIN = 457200  # bottom margin value for Narrow margins as interpreted by python-docx
        LEFT_MARGIN = 457200  # left margin value for Narrow margins as interpreted by python-docx
        RIGHT_MARGIN = 457200  # right margin value for Narrow margins as interpreted by python-docx
        top, bottom, left, right = margins
        return top == TOP_MARGIN and bottom == BOTTOM_MARGIN and left == LEFT_MARGIN and right == RIGHT_MARGIN

    # obtain number of sections that do not contain Narrow margins and also convert them into strings
    no_Narrow = [str(i + 1) for i, section in enumerate(page_margins) if not Narrow_present(section)]

    # check if the length of the lists no_Narrow and page_margins are equal, if so this means that all pages do not have Narrow margins
    if len(no_Narrow) == len(page_margins):
        # print this if the length of the lists no_Narrow and page_margins are equal, meaning that all pages do not have Narrow margins
        st.error("❌ Whole document does not have Narrow margins.")

    # check if the list no_Narrow contains one element, if so this means that one section contains non-Narrow margins
    elif len(no_Narrow) == 1:
        # print this if the list no_Narrow contains one element, meaning that one section contains non-Narrow margins
        st.error(f"❌ Section {no_Narrow[0]} does not have Narrow margins.")

    # check if the list no_Narrow is not empty and if the length of the list page_margins is greater than the length of the list no_Narrow
    elif 0 < len(no_Narrow) < len(page_margins):
        st.error(f"❌ Sections {', '.join(no_Narrow)} do not have Narrow margins.")


# function to run selected programs(s)
def run_program():
    if cb2:
        font_name()
    if cb3:
        font_size()
    if cb4:
        font_colour()
    if cb5:
        font_style()
    if cb6:
        para_align()
    if cb7:
        spacing()
    if cb8:
        ref_count()
    if cb9:
        header_footer()
    if cb10:
        page_margin()


# configure sidebar text and widgets
st.sidebar.title("**Format Check v1.1**")
# select a document using file uploader
uploaded_file = st.sidebar.file_uploader("Choose a Word document", type='.docx', key=1)
cb1 = st.sidebar.checkbox('All', key=2)
# if 'All Formatting Items' is checked disable every other checkbox
if cb1:
    cb2 = st.sidebar.checkbox('Font Name', value=cb1, disabled=True, key=3)
    cb3 = st.sidebar.checkbox('Font Size', value=cb1, disabled=True, key=4)
    cb4 = st.sidebar.checkbox('Font Colour', value=cb1, disabled=True, key=5)
    cb5 = st.sidebar.checkbox('Font Style', value=cb1, disabled=True, key=6)
    cb6 = st.sidebar.checkbox('Paragraph Alignment', value=cb1, disabled=True, key=7)
    cb7 = st.sidebar.checkbox('Spacing', value=cb1, disabled=True, key=8)
    cb8 = st.sidebar.checkbox('Reference Counter', value=cb1, disabled=True, key=9)
    cb9 = st.sidebar.checkbox('Header and Footer', value=cb1, disabled=True, key=10)
    cb10 = st.sidebar.checkbox('Page Margins', value=cb1, disabled=True, key=11)
else:
    cb2 = st.sidebar.checkbox('Font Name', key=12)
    cb3 = st.sidebar.checkbox('Font Size', key=13)
    cb4 = st.sidebar.checkbox('Font Colour', key=14)
    cb5 = st.sidebar.checkbox('Font Style', key=15)
    cb6 = st.sidebar.checkbox('Paragraph Alignment', key=16)
    cb7 = st.sidebar.checkbox('Spacing', key=17)
    cb8 = st.sidebar.checkbox('Reference Counter', key=18)
    cb9 = st.sidebar.checkbox('Header and Footer', key=19)
    cb10 = st.sidebar.checkbox('Page Margins', key=20)

start_btn = st.sidebar.button('Start', on_click=run_program, key=21)
st.sidebar.text("")
st.sidebar.text("")

# style Start button
button_style = """
        <style>
        .stButton > button {
            color: black;
            font-size: 20px;
            background-color: #FFFFFF;
            border: 1px solid black;
            width: 305px;
            height: 50px;
        }
        </style>
        """
st.markdown(button_style, unsafe_allow_html=True)


# Hide hamburger menu and 'Made with Streamlit' footer
hide_streamlit_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            </style>
            """
st.markdown(hide_streamlit_style, unsafe_allow_html=True)
